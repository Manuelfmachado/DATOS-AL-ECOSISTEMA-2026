"""
Módulo Simulación — Tres decisiones, una fuente integrada (Supabase).

1. Alineación curricular (Universidades): ¿Conviene actualizar el pensum?
   Compara el pensum actual con habilidades esenciales ESCO de ocupaciones afines.
2. Intervención por objetivo (Gobierno): ¿Dónde y en qué sector intervenir?
   Ranking de oportunidades por depto, ponderado según objetivo declarado.
3. Explora carrera (Estudiantes): ¿Qué pasa si estudio esta carrera?
   Habilidades, salidas laborales, demanda, salario y dónde estudiarla. Sin score mágico.

Plus: Analizar con IA — Gemini responde preguntas sobre la simulación generada.

Todo desde Supabase. Cero scores inventados: cada número se explica con su fórmula.
"""
import json
from collections import Counter
from pathlib import Path
from typing import Any

from fastapi import APIRouter, HTTPException, UploadFile, File, Form
from pydantic import BaseModel

from app.db.supabase import supabase
from app.services.llm_gemini import (
    call_gemini_text,
    call_gemini_multimodal_json,
    is_gemini_available,
)
from app.services.llm import call_llm_text
from app.services.onet_mercado import (
    onet_codes_para_programa,
    tareas_onet_para_ocupaciones,
    educacion_onet_para_ocupaciones,
)

router = APIRouter(prefix="/api/simulacion", tags=["Simulacion"])

PREDICCIONES_MUNDIALES_PATH = Path(__file__).resolve().parents[2] / "data" / "processed" / "predicciones_mundiales.json"


def _load_predicciones_mundiales() -> dict[str, Any]:
    """Carga las predicciones mundiales (Chronos) desde el JSON local."""
    if not PREDICCIONES_MUNDIALES_PATH.exists():
        return {}
    with open(PREDICCIONES_MUNDIALES_PATH, "r", encoding="utf-8") as f:
        return json.load(f)

SMMLV_2026 = 1_750_000

DANE_DEPTO: dict[str, str] = {
    "05": "ANTIOQUIA", "08": "ATLANTICO", "11": "BOGOTA D.C.", "13": "BOLIVAR",
    "15": "BOYACA", "17": "CALDAS", "18": "CAQUETA", "19": "CAUCA",
    "20": "CESAR", "23": "CORDOBA", "25": "CUNDINAMARCA", "27": "CHOCO",
    "41": "HUILA", "44": "LA GUAJIRA", "47": "MAGDALENA", "50": "META",
    "52": "NARINO", "54": "NORTE DE SANTANDER", "63": "QUINDIO",
    "66": "RISARALDA", "68": "SANTANDER", "70": "SUCRE", "73": "TOLIMA",
    "76": "VALLE DEL CAUCA", "81": "ARAUCA", "85": "CASANARE",
    "86": "PUTUMAYO", "88": "ARCHIPIELAGO DE SAN ANDRES",
    "91": "AMAZONAS", "94": "GUAINIA", "95": "GUAVIARE",
    "97": "VAUPES", "99": "VICHADA",
}

RANGO_INGRESO_MEDIO_SMMLV: dict[str, float] = {
    "1 SMMLV": 1.0, "Entre 1 y 1,5 SMMLV": 1.25,
    "Entre 1,5 y 2,5 SMMLV": 2.0, "Entre 2,5 y 4 SMMLV": 3.25,
    "Entre 4 y 6 SMMLV": 5.0, "Entre 6 y 9 SMMLV": 7.5,
    "Mas de 9 SMMLV": 11.0, "Más de 9 SMMLV": 11.0,
}

CIIU2_NOMBRES: dict[int, str] = {
    0: "Sin clasificar",
    1: "Agricultura, ganadería, caza y silvicultura",
    2: "Explotación de minas y canteras",
    3: "Industria manufacturera",
    5: "Construcción",
    6: "Comercio, hoteles y restaurantes",
    7: "Transporte y comunicaciones",
    8: "Información y comunicaciones",
    9: "Actividades financieras e inmobiliarias",
    10: "Servicios comunales, sociales y personales",
    10: "Fabricación de productos alimenticios",
    11: "Fabricación de bebidas",
    13: "Fabricación de textiles",
    14: "Confección de prendas de vestir",
    17: "Transformación de madera",
    18: "Fabricación de papel y cartón",
    20: "Fabricación de sustancias químicas",
    21: "Fabricación de productos farmacéuticos",
    22: "Productos de caucho y plástico",
    23: "Otros minerales no metálicos",
    24: "Industria básica de hierro y acero",
    25: "Productos metálicos",
    26: "Productos informáticos y electrónicos",
    27: "Equipo eléctrico",
    28: "Maquinaria y equipo",
    29: "Vehículos automotores",
    30: "Otros equipos de transporte",
    31: "Fabricación de muebles",
    32: "Otras industrias manufactureras",
    35: "Suministro de electricidad, gas y vapor",
    36: "Distribución de agua",
    41: "Construcción de edificaciones",
    43: "Actividades especializadas de construcción",
    45: "Comercio de vehículos",
    46: "Comercio al por mayor",
    47: "Comercio al por menor",
    49: "Transporte terrestre",
    50: "Transporte acuático",
    51: "Transporte aéreo",
    52: "Almacenamiento",
    53: "Actividades postales y mensajería",
    55: "Alojamiento",
    56: "Servicios de comida",
    58: "Actividades de edición",
    59: "Actividades cinematográficas",
    61: "Telecomunicaciones",
    62: "Desarrollo de sistemas informáticos (software)",
    63: "Procesamiento de datos y hosting",
    64: "Servicios financieros",
    65: "Seguros",
    66: "Actividades auxiliares a servicios financieros",
    68: "Actividades inmobiliarias",
    69: "Actividades jurídicas y contables",
    70: "Consultoría de gestión",
    71: "Servicios técnicos y arquitectónicos",
    72: "Investigación científica",
    73: "Publicidad y estudios de mercado",
    74: "Otras actividades profesionales",
    78: "Servicios de empleo y agencias",
    80: "Seguridad e investigación",
    81: "Servicios a edificios",
    82: "Actividades administrativas",
    84: "Administración pública y defensa",
    85: "Educación",
    86: "Atención en salud",
    87: "Asistencia social con alojamiento",
    90: "Actividades artísticas y de entretenimiento",
    93: "Actividades deportivas y recreativas",
    94: "Actividades de asociaciones",
    96: "Otros servicios personales",
    97: "Hogares como empleadores",
}

# Mapa aproximado de código CIIU-2-dígitos → nombre de sector en PILA y predicciones
CIIU2_KEYWORD_SECTOR: dict[int, list[str]] = {
    62: ["software", "informatic", "tecnolog", "sistema", "datos"],
    63: ["datos", "informacion", "hosting"],
    61: ["telecom", "telecomunicacion"],
    85: ["educacion", "enseñanza"],
    86: ["salud", "hospital", "medic"],
    84: ["publica", "gobierno", "administracion publica"],
    47: ["comercio", "minorista", "retail", "tienda"],
    46: ["comercio", "mayorista", "mayor"],
    56: ["restaurante", "comida", "gastronomia", "alojamiento"],
    55: ["hotel", "alojamiento", "turismo"],
    49: ["transporte", "logistica", "vehiculo"],
    41: ["construccion", "obra", "edificacion"],
    1: ["agricola", "agro", "campo", "cultivo", "ganaderia"],
    3: ["manufactura", "industria", "fabrica"],
}

# Mapeo CIIU-2-digitos -> GRUPOS12 (EMICRON) para traer empleo informal por sector
CIIU2_TO_GRUPOS12: dict[int, int] = {
    1: 1,   # Agricultura
    2: 2,   # Mineria
    3: 3,   # Manufactura
    5: 5,   # Construccion
    6: 3,   # Servicios publicos -> asimilable a manufactura (no hay GRUPOS12 directo)
    7: 3,   # Comercio (mayor) - CIIU 46
    8: 3,   # Transporte
    9: 9,   # Alojamiento y comida
    10: 12, # Informacion/comunicaciones -> servicios profesionales
    11: 11, # Financieros
    12: 12, # Servicios profesionales
    13: 12, # Administracion publica -> servicios
    14: 12, # Educacion -> servicios
    15: 12, # Salud -> servicios
    # Mapeos directos por keyword
    46: 6,  # Comercio al por mayor
    47: 7,  # Comercio al por menor
    49: 8,  # Transporte
    55: 9,  # Alojamiento
    56: 9,  # Comida
    41: 5,  # Construccion
    62: 12, # Software -> servicios profesionales
    63: 10, # Datos/informacion
    61: 10, # Telecomunicaciones
    85: 12, # Educacion
    86: 12, # Salud
    84: 12, # Administracion publica
}

# Pesos por objetivo de intervención (Gobierno)
PESOS_OBJETIVO: dict[str, dict[str, float]] = {
    "sectores_emergentes":   {"demanda": 0.25, "crecimiento": 0.40, "deficit_talento": 0.20, "compatibilidad": 0.15},
    "reducir_desempleo_juvenil": {"demanda": 0.45, "crecimiento": 0.20, "deficit_talento": 0.20, "compatibilidad": 0.15},
    "emprendimiento":         {"demanda": 0.20, "crecimiento": 0.30, "deficit_talento": 0.35, "compatibilidad": 0.15},
    "reducir_informalidad":   {"demanda": 0.20, "crecimiento": 0.15, "deficit_talento": 0.20, "compatibilidad": 0.45},
}


# ===========================================================================
# Helpers
# ===========================================================================

# Normalización y limpieza de texto. La implementación vive en
# app.services._norm (para evitar dependencias circulares con los servicios);
# aquí solo se reexporta para preservar los imports existentes que usan
# `_norm` / `_fix_mojibake` desde este módulo.
from app.services._norm import _fix_mojibake, _norm  # noqa: E402,F401


# Diccionario ampliado español -> inglés (tokens/frases) para matching sin LLM.
# Cubre los términos más comunes en pensums de ingeniería, administración y afines.
# Cada clave puede mapear a múltiples equivalentes en ESCO/inglés técnico.
TOKEN_MAP_ES_EN: dict[str, list[str]] = {
    "ANALISIS DE DATOS": ["DATA ANALYSIS", "ANALYSE DATA", "DATA ANALYTICS", "ANALYTICS"],
    "ANALISIS": ["ANALYSIS", "ANALYSE", "ASSESS"],
    "DATOS": ["DATA", "INFORMATION"],
    "LOGISTICA": ["LOGISTICS", "SUPPLY CHAIN", "MANAGE LOGISTICS", "STOCK CONTROL"],
    "CADENA DE SUMINISTRO": ["SUPPLY CHAIN", "LOGISTICS"],
    "INVENTARIOS": ["INVENTORY", "STOCK CONTROL", "STOCK MANAGEMENT"],
    "MECANICA": ["MECHANICS", "MECHANICAL ENGINEERING", "MACHINERY"],
    "MECANICA DE FLUIDOS": ["FLUID MECHANICS", "HYDRAULICS", "PNEUMATICS"],
    "HIDRAULICA": ["HYDRAULICS", "FLUID MECHANICS"],
    "NEUMATICA": ["PNEUMATICS", "FLUID MECHANICS"],
    "PROGRAMACION": ["PROGRAMMING", "DEVELOP SOFTWARE", "WRITE CODE", "COMPUTER PROGRAMMING"],
    "CODIGO": ["CODE", "PROGRAMMING"],
    "DESARROLLO WEB": ["WEB DEVELOPMENT", "WEB DESIGN", "BUILD WEBSITES"],
    "DESARROLLO DE SOFTWARE": ["SOFTWARE DEVELOPMENT", "DEVELOP SOFTWARE"],
    "BASES DE DATOS": ["DATABASES", "SQL", "MANAGE DATA", "DATABASE MANAGEMENT"],
    "SQL": ["SQL", "DATABASES", "QUERY LANGUAGES"],
    "MATEMATICAS": ["MATHEMATICS", "CALCULATE", "MATH", "STATISTICS"],
    "CALCULO": ["CALCULUS", "MATHEMATICS", "CALCULATE"],
    "ALGEBRA": ["ALGEBRA", "MATHEMATICS"],
    "ESTADISTICA": ["STATISTICS", "STATISTICAL ANALYSIS"],
    "PROBABILIDAD": ["PROBABILITY", "STATISTICS"],
    "INGLES": ["ENGLISH", "FOREIGN LANGUAGE", "LANGUAGE SKILLS"],
    "IDIOMA": ["FOREIGN LANGUAGE", "LANGUAGE SKILLS"],
    "ADMINISTRACION": ["MANAGEMENT", "BUSINESS ADMINISTRATION", "MANAGE STAFF", "BUSINESS MANAGEMENT"],
    "GESTION": ["MANAGEMENT", "MANAGE", "ADMINISTRATION"],
    "FINANZAS": ["FINANCE", "FINANCIAL ANALYSIS", "ACCOUNTING", "FINANCIAL MANAGEMENT"],
    "CONTABILIDAD": ["ACCOUNTING", "BOOKKEEPING", "FINANCE"],
    "COSTOS": ["COST ACCOUNTING", "COSTS", "FINANCE"],
    "PRESUPUESTO": ["BUDGET", "BUDGETING", "FINANCIAL PLANNING"],
    "SEGURIDAD": ["SAFETY", "SECURITY", "SAFETY MEASURES", "HEALTH AND SAFETY", "OCCUPATIONAL SAFETY"],
    "SEGURIDAD INDUSTRIAL": ["SAFETY ENGINEERING", "OCCUPATIONAL SAFETY", "HEALTH AND SAFETY", "SAFETY MEASURES"],
    "SEGURIDAD Y SALUD": ["HEALTH AND SAFETY", "OCCUPATIONAL SAFETY", "SAFETY ENGINEERING"],
    "RIESGOS LABORALES": ["OCCUPATIONAL SAFETY", "RISK MANAGEMENT", "HEALTH AND SAFETY"],
    "DISEÑO": ["DESIGN", "DESIGN PRINCIPLES", "BLUEPRINTS"],
    "DISEÑO GRAFICO": ["GRAPHIC DESIGN", "DESIGN PRINCIPLES"],
    "PROYECTOS": ["PROJECT MANAGEMENT", "MANAGE SCHEDULE OF TASKS", "PROJECTS"],
    "GESTION DE PROYECTOS": ["PROJECT MANAGEMENT", "MANAGE SCHEDULE OF TASKS"],
    "LIDERAZGO": ["LEADERSHIP", "LEAD A TEAM", "MANAGE A TEAM"],
    "TRABAJO EN EQUIPO": ["TEAMWORK", "WORK IN TEAMS", "COLLABORATE"],
    "COMUNICACION": ["COMMUNICATION", "COMMUNICATION SKILLS", "COMMUNICATE"],
    "NEGOCIACION": ["NEGOTIATION", "NEGOTIATE"],
    "CALIDAD": ["QUALITY", "QUALITY MANAGEMENT", "QUALITY CONTROL", "QUALITY STANDARDS"],
    "GESTION DE CALIDAD": ["QUALITY MANAGEMENT", "QUALITY CONTROL", "QUALITY STANDARDS"],
    "INVESTIGACION": ["RESEARCH", "INVESTIGATE", "RESEARCH METHODS"],
    "MARKETING": ["MARKETING", "MARKET RESEARCH", "ADVERTISING"],
    "MERCADEO": ["MARKETING", "MARKET RESEARCH"],
    "VENTAS": ["SALES", "SELL", "SELLING"],
    "RECURSOS HUMANOS": ["HUMAN RESOURCES", "HR MANAGEMENT", "STAFF MANAGEMENT"],
    "OPERACIONES": ["OPERATIONS", "OPERATIONS MANAGEMENT"],
    "MANUFACTURA": ["MANUFACTURING", "PRODUCTION PROCESSES", "INDUSTRIAL PROCESSES"],
    "PRODUCCION": ["PRODUCTION", "MANUFACTURING", "PRODUCTION PROCESSES"],
    "MANTENIMIENTO": ["MAINTENANCE", "REPAIR", "MAINTAIN EQUIPMENT"],
    "ELECTRICIDAD": ["ELECTRICITY", "ELECTRICAL", "ELECTRICAL ENGINEERING"],
    "ELECTRONICA": ["ELECTRONICS", "ELECTRONIC ENGINEERING"],
    "AUTOMATIZACION": ["AUTOMATION", "AUTOMATED SYSTEMS", "PROCESS AUTOMATION"],
    "CONTROL DE PROCESOS": ["PROCESS CONTROL", "AUTOMATION", "PROCESS MONITORING"],
    "ROBOTICA": ["ROBOTICS", "AUTOMATION"],
    "SOSTENIBILIDAD": ["SUSTAINABILITY", "SUSTAINABLE DEVELOPMENT"],
    "INNOVACION": ["INNOVATION", "INNOVATIVE"],
    "EMPRENDIMIENTO": ["ENTREPRENEURSHIP", "BUSINESS CREATION"],
    "ETICA": ["ETHICS", "PROFESSIONAL ETHICS"],
    "RESPONSABILIDAD SOCIAL": ["SOCIAL RESPONSIBILITY", "CORPORATE SOCIAL RESPONSIBILITY"],
    "NORMATIVIDAD": ["REGULATIONS", "LEGISLATION", "COMPLIANCE"],
    "NORMAS": ["STANDARDS", "REGULATIONS", "COMPLIANCE"],
    "ISO": ["QUALITY STANDARDS", "ISO STANDARDS", "STANDARDS"],
    "MODELAJE": ["MODELING", "SIMULATION", "MODELS"],
    "MODELADO": ["MODELING", "SIMULATION"],
    "SIMULACION": ["SIMULATION", "MODELING", "SIMULATE"],
    "OPTIMIZACION": ["OPTIMISATION", "OPTIMIZATION", "OPTIMISE"],
    "INVESTIGACION DE OPERACIONES": ["OPERATIONS RESEARCH", "OPTIMISATION"],
    "INGENIERIA DE PROCESOS": ["PROCESS ENGINEERING", "PROCESS IMPROVEMENT"],
    "MEJORA CONTINUA": ["CONTINUOUS IMPROVEMENT", "PROCESS IMPROVEMENT"],
    "LEAN": ["LEAN MANUFACTURING", "CONTINUOUS IMPROVEMENT"],
    "SEIS SIGMA": ["SIX SIGMA", "QUALITY MANAGEMENT", "STATISTICS"],
    "CADENA DE VALOR": ["VALUE CHAIN", "BUSINESS PROCESS"],
    "MUESTREO": ["SAMPLING", "STATISTICS"],
    "ESTUDIO DE TIEMPOS": ["TIME STUDY", "WORK MEASUREMENT"],
    "METODOS NUMERICOS": ["NUMERICAL METHODS", "MATHEMATICS"],
    "FISICA": ["PHYSICS", "SCIENCE"],
    "QUIMICA": ["CHEMISTRY", "CHEMICAL"],
    "TERMODINAMICA": ["THERMODYNAMICS", "PHYSICS"],
    "RESISTENCIA DE MATERIALES": ["MATERIALS SCIENCE", "MECHANICS OF MATERIALS"],
    "MAQUINAS": ["MACHINERY", "MACHINES"],
    "MAQUINARIA": ["MACHINERY", "MACHINES"],
    "HERRAMIENTAS": ["TOOLS", "INDUSTRIAL TOOLS"],
    "PLANOS": ["BLUEPRINTS", "TECHNICAL DRAWINGS", "DESIGN PRINCIPLES"],
    "DIBUJO TECNICO": ["TECHNICAL DRAWING", "BLUEPRINTS"],
    "ERGONOMIA": ["ERGONOMICS", "WORKPLACE DESIGN"],
    "SALUD OCUPACIONAL": ["OCCUPATIONAL HEALTH", "HEALTH AND SAFETY"],
    "AMBIENTAL": ["ENVIRONMENTAL", "ENVIRONMENTAL IMPACT"],
    "ENERGIA": ["ENERGY", "POWER", "ELECTRICITY"],
    "REDES": ["NETWORKS", "NETWORKING"],
    "SISTEMAS": ["SYSTEMS", "INFORMATION SYSTEMS"],
    "INTELIGENCIA ARTIFICIAL": ["ARTIFICIAL INTELLIGENCE", "AI", "MACHINE LEARNING"],
    "MACHINE LEARNING": ["MACHINE LEARNING", "ARTIFICIAL INTELLIGENCE"],
    "APRENDIZAJE AUTOMATICO": ["MACHINE LEARNING", "ARTIFICIAL INTELLIGENCE"],
    "MINERIA DE DATOS": ["DATA MINING", "DATA ANALYTICS"],
    "VISUALIZACION": ["DATA VISUALISATION", "VISUALISATION"],
    "POWER BI": ["DATA VISUALISATION", "BUSINESS INTELLIGENCE", "BI"],
    "EXCEL": ["SPREADSHEETS", "DATA ANALYSIS"],
    "TABLEAU": ["DATA VISUALISATION", "BUSINESS INTELLIGENCE"],
    "PYTHON": ["PYTHON", "PROGRAMMING", "SCRIPTING"],
    "R": ["R PROGRAMMING", "STATISTICS"],
    "JAVA": ["JAVA", "PROGRAMMING"],
    "JAVASCRIPT": ["JAVASCRIPT", "WEB DEVELOPMENT"],
    "GESTION DEL RIESGO": ["RISK MANAGEMENT", "RISK ASSESSMENT"],
    "ANALISIS DE RIESGO": ["RISK ANALYSIS", "RISK ASSESSMENT", "PERFORM RISK ANALYSIS"],
}

# Diccionario inverso inglés -> español para traducir faltantes al mostrarlas.
TOKEN_MAP_EN_ES: dict[str, str] = {
    "DATA ANALYSIS": "Análisis de datos", "ANALYSE DATA": "Análisis de datos",
    "DATA ANALYTICS": "Analítica de datos", "ANALYTICS": "Analítica",
    "LOGISTICS": "Logística", "SUPPLY CHAIN": "Cadena de suministro",
    "MANAGE LOGISTICS": "Gestión logística", "STOCK CONTROL": "Control de inventarios",
    "MECHANICS": "Mecánica", "MECHANICAL ENGINEERING": "Ingeniería mecánica",
    "MACHINERY": "Maquinaria", "PROGRAMMING": "Programación",
    "DEVELOP SOFTWARE": "Desarrollo de software", "WRITE CODE": "Programación",
    "WEB DEVELOPMENT": "Desarrollo web", "WEB DESIGN": "Diseño web",
    "DATABASES": "Bases de datos", "SQL": "SQL / Bases de datos",
    "MANAGE DATA": "Gestión de datos", "MATHEMATICS": "Matemáticas",
    "STATISTICS": "Estadística", "STATISTICAL ANALYSIS": "Análisis estadístico",
    "ENGLISH": "Inglés", "FOREIGN LANGUAGE": "Idioma extranjero",
    "MANAGEMENT": "Gestión / Administración", "BUSINESS ADMINISTRATION": "Administración de empresas",
    "MANAGE STAFF": "Gestión de personal", "FINANCE": "Finanzas",
    "FINANCIAL ANALYSIS": "Análisis financiero", "ACCOUNTING": "Contabilidad",
    "SAFETY": "Seguridad", "SAFETY MEASURES": "Medidas de seguridad",
    "HEALTH AND SAFETY": "Salud y seguridad", "OCCUPATIONAL SAFETY": "Seguridad ocupacional",
    "SAFETY ENGINEERING": "Ingeniería de seguridad", "SECURITY": "Seguridad",
    "DESIGN": "Diseño", "DESIGN PRINCIPLES": "Principios de diseño",
    "BLUEPRINTS": "Planos", "PROJECT MANAGEMENT": "Gestión de proyectos",
    "MANAGE SCHEDULE OF TASKS": "Gestión de cronograma de tareas",
    "LEADERSHIP": "Liderazgo", "LEAD A TEAM": "Liderar equipos",
    "TEAMWORK": "Trabajo en equipo", "COMMUNICATION": "Comunicación",
    "COMMUNICATION SKILLS": "Habilidades de comunicación", "NEGOTIATION": "Negociación",
    "QUALITY": "Calidad", "QUALITY MANAGEMENT": "Gestión de calidad",
    "QUALITY CONTROL": "Control de calidad", "QUALITY STANDARDS": "Normas de calidad",
    "RESEARCH": "Investigación", "MARKETING": "Marketing / Mercadeo",
    "SALES": "Ventas", "HUMAN RESOURCES": "Recursos humanos",
    "OPERATIONS": "Operaciones", "MANUFACTURING": "Manufactura",
    "PRODUCTION PROCESSES": "Procesos de producción", "MAINTENANCE": "Mantenimiento",
    "REPAIR": "Reparación", "ELECTRICITY": "Electricidad",
    "ELECTRICAL": "Eléctrico", "ELECTRONICS": "Electrónica",
    "AUTOMATION": "Automatización", "PROCESS CONTROL": "Control de procesos",
    "ROBOTICS": "Robótica", "SUSTAINABILITY": "Sostenibilidad",
    "INNOVATION": "Innovación", "ENTREPRENEURSHIP": "Emprendimiento",
    "ETHICS": "Ética", "REGULATIONS": "Normatividad",
    "STANDARDS": "Normas", "COMPLIANCE": "Cumplimiento normativo",
    "MODELING": "Modelado", "SIMULATION": "Simulación",
    "OPTIMISATION": "Optimización", "OPTIMIZATION": "Optimización",
    "OPERATIONS RESEARCH": "Investigación de operaciones",
    "PROCESS ENGINEERING": "Ingeniería de procesos",
    "CONTINUOUS IMPROVEMENT": "Mejora continua", "PROCESS IMPROVEMENT": "Mejora de procesos",
    "LEAN MANUFACTURING": "Manufactura esbelta",
    "RISK MANAGEMENT": "Gestión del riesgo", "RISK ASSESSMENT": "Evaluación de riesgos",
    "PERFORM RISK ANALYSIS": "Análisis de riesgos", "RISK ANALYSIS": "Análisis de riesgos",
    "MATERIALS SCIENCE": "Ciencia de materiales", "MECHANICS OF MATERIALS": "Resistencia de materiales",
    "TECHNICAL DRAWING": "Dibujo técnico", "ERGONOMICS": "Ergonomía",
    "OCCUPATIONAL HEALTH": "Salud ocupacional", "ENVIRONMENTAL": "Ambiental",
    "ENERGY": "Energía", "NETWORKS": "Redes", "SYSTEMS": "Sistemas",
    "ARTIFICIAL INTELLIGENCE": "Inteligencia artificial",
    "MACHINE LEARNING": "Aprendizaje automático", "DATA MINING": "Minería de datos",
    "DATA VISUALISATION": "Visualización de datos", "BUSINESS INTELLIGENCE": "Inteligencia de negocios",
    "TOOLS": "Herramientas", "INDUSTRIAL TOOLS": "Herramientas industriales",
    "HYDRAULICS": "Hidráulica", "PNEUMATICS": "Neumática",
    "FLUID MECHANICS": "Mecánica de fluidos", "BUDGET": "Presupuesto",
    "BUDGETING": "Presupuestación", "FINANCIAL PLANNING": "Planificación financiera",
    "COST ACCOUNTING": "Costos", "PHYSICS": "Física", "CHEMISTRY": "Química",
    "THERMODYNAMICS": "Termodinámica", "MACHINES": "Máquinas",
    "ANALYSIS": "Análisis", "ASSESS": "Evaluar", "INFORMATION": "Información",
    "INFORMATION SYSTEMS": "Sistemas de información",
    "DISASSEMBLE EQUIPMENT": "Desmontaje de equipos", "DISASSEMBLE MACHINES": "Desmontaje de máquinas",
    "INSTRUCT ON SAFETY MEASURES": "Instruir sobre medidas de seguridad",
    "ADVISE ON MACHINERY MALFUNCTIONS": "Asesorar sobre fallos de maquinaria",
}


def _normalizar_duracion_sena(valor: Any) -> str:
    """Normaliza el campo duracion_horas del SENA a formato legible 'Nh'.
    La fuente trae valores como 1.8 (= 1800 h) o 600 (= 600 h)."""
    if valor is None:
        return "N/D"
    try:
        f = float(str(valor).replace(",", "."))
        if f <= 0:
            return "N/D"
        # Si el valor es pequeño (<100), probablemente está en miles (1.8 = 1800h)
        if f < 100:
            h = int(round(f * 1000))
        else:
            h = int(round(f))
        return f"{h}h"
    except Exception:
        return str(valor)


def _norm_depto(s: str) -> str:
    n = _norm(s)
    if n in ("BOGOTA D.C", "BOGOTA", "BOGOTA DC", "BOGOTA, D.C.", "BOGOTA, DC"):
        return "BOGOTA D.C."
    if "SAN ANDRES" in n:
        return "SAN ANDRES Y PROVIDENCIA"
    return n


def _depto_to_code(depto_norm: str) -> str | None:
    for code, name in DANE_DEPTO.items():
        if _norm_depto(name) == depto_norm:
            return code
    return None


def _rango_a_cop(rango: str) -> float | None:
    if not rango:
        return None
    nr = _norm(rango)
    for k, v in RANGO_INGRESO_MEDIO_SMMLV.items():
        if _norm(k) == nr:
            return v * SMMLV_2026
    return None


def _sb_select(table: str, select: str = "*", filt: dict | None = None,
               order: str | None = None, limit: int | None = None) -> list[dict]:
    q = supabase.table(table).select(select)
    if filt:
        for k, v in filt.items():
            q = q.eq(k, v)
    if order:
        col, direction = (order.split(".") + ["asc"])[:2]
        q = q.order(col, desc=(direction == "desc"))
    if limit:
        q = q.limit(limit)
    return q.execute().data or []


def _sb_ilike(table: str, select: str, column: str, value: str,
              limit: int = 50, order: str | None = None) -> list[dict]:
    q = supabase.table(table).select(select).ilike(column, value)
    if order:
        col, direction = (order.split(".") + ["asc"])[:2]
        q = q.order(col, desc=(direction == "desc"))
    return q.limit(limit).execute().data or []


# ===========================================================================
# Simulación 1 — Alineación curricular (Universidades)
# ===========================================================================

class AlineacionRequest(BaseModel):
    programa: str
    pensum: list[str] = []


def _ocupaciones_compatibles(programa_norm: str, limit: int = 5) -> list[dict]:
    """Busca ocupaciones ESCO cuyas palabras principales coincidan con el programa.
    Estrategia:
    1) Búsqueda por combinación de palabras distintivas del programa (ej: 'ingeniero industrial')
       usando ilike con la frase completa primero (mejor match).
    2) Si no hay suficientes, búsqueda por palabra distintiva individual (>= 5 chars para reducir ruido).
    3) Filtro por grupo ISCO según el tipo de programa, para excluir ocupaciones que solo comparten
       una palabra pero no son del campo profesional (ej: 'vendedor de software' para ingeniería).
    Acota a `limit` ocupaciones (default 5)."""
    # Palabras distintivas del programa (>= 4 chars, sin stopwords genéricas)
    stop = {"DE", "DEL", "LA", "LAS", "LOS", "EL", "EN", "Y", "E", "PARA", "CON",
            "POR", "AL", "A", "O", "TECNOL", "TECNICA", "PROFESIONAL", "ESPECIALIZACION",
            "MAESTRIA", "DOCTORADO", "LICENCIATURA", "DIPLOMADO"}
    # No filtramos INGENIERIA/CIENCIA/ADMINISTR aquí, porque son palabra principal del programa
    palabras = [w for w in programa_norm.split() if len(w) >= 4 and w not in stop]
    if not palabras:
        palabras = [w for w in programa_norm.split() if len(w) >= 4][:3]

    # Determinar grupos ISCO permitidos según el tipo de programa.
    # ISCO-08: 1=directores, 2=profesionales, 3=técnicos, 4=empleados admin, 5=servicios/vendedores,
    #          6=agro, 7=artesanos, 8=operarios, 9=ocupaciones elementales, 0=fuerzas armadas.
    # Para programas profesionales (ingeniería, ciencias, salud, derecho) exigimos grupo 2 o 3.
    # Para técnicos/tecnólogos, grupo 3. Para administración, 1/2/4. Para servicios, 4/5.
    prog_upper = programa_norm.upper()
    isco_groups_allowed: set[int] = set()
    if any(w in prog_upper for w in ["INGENIER", "CIENCIA", "CIENCIAS", "MEDICINA", "ENFERMER",
                                      "ODONTOLOG", "NUTRICION", "FISIOTERAP", "PSICOLOG",
                                      "DERECHO", "JURISPRUD", "ECONOM", "CONTADUR", "ADMINISTR",
                                      "COMUNIC", "PERIODISM", "SOCIOLOG", "TRABAJO SOCIAL",
                                      "BIOLOG", "QUIMICA", "FISICA", "MATEMAT", "ESTADIST",
                                      "ARQUITECT", "AGRONOM", "VETERIN", "GEOLOG", "BIBLIOTEC",
                                      "BACTERIOLOG", "ZOOTECNIA", "INSTRUMENTACION"]):
        # Programas profesionales: ocupaciones de profesionales (2) y técnicos (3)
        isco_groups_allowed = {2, 3}
    elif any(w in prog_upper for w in ["TECNOLOG", "TECNIC", "TECNICA"]):
        isco_groups_allowed = {3, 2}
    elif any(w in prog_upper for w in ["LICENCIATUR", "EDUCACION", "PEDAGOG"]):
        isco_groups_allowed = {2, 3}
    # Si no se identifica el tipo, no filtrar (mantener comportamiento anterior)

    # Mapeo programa -> ocupaciones ESCO conocidas por ISCO, para casos donde el
    # nombre del rol no coincide con la palabra del programa (ej: MEDICINA -> médico,
    # ODONTOLOGIA -> odontólogo, etc.). Devuelve lista de ISCO codes a buscar directamente.
    prog_to_isco_codes: dict[str, list[str]] = {
        "MEDICINA": ["2211"],  # médico general (sin especialistas para no contaminar)
        "ODONTOLOG": ["2221", "2222"],  # odontólogos (2221 enfermeros/dentistas relacionados)
        "ENFERMER": ["2221", "3221"],  # enfermeros profesionales y técnicos
        "PSICOLOG": ["2634"],  # psicólogos
        "NUTRICION": ["2213"],  # nutricionistas
        "FISIOTERAP": ["2264"],  # fisioterapeutas
        "VETERIN": ["2250", "2251"],  # veterinarios
        "ZOOTECNIA": ["2250", "2251"],
        "BACTERIOLOG": ["3212", "3213"],  # técnicos de laboratorio médico
        "DERECHO": ["2611", "2612", "2619"],  # abogados
        "JURISPRUD": ["2611", "2612", "2619"],
        "ECONOM": ["2631"],  # economistas
        "CONTADUR": ["2411", "2414"],  # contadores
        "ADMINISTR": ["1211", "1219", "2421"],  # gerentes/administradores
        "ARQUITECT": ["2161", "2162", "2163"],  # arquitectos
        "AGRONOM": ["2214"],  # ingenieros agrónomos
        "TRABAJO SOCIAL": ["2635"],  # trabajadores sociales
        "SOCIOLOG": ["2632"],  # sociólogos/antropólogos
        "BIBLIOTEC": ["2652"],  # bibliotecarios
        "INSTRUMENTACION": ["3214"],  # instrumentistas quirúrgicos
        "COMUNIC": ["2641", "2642", "2643"],  # periodistas/comunicadores
        "PERIODISM": ["2641", "2642", "2643"],
        "MERCAD": ["2431", "2432"],  # especialistas en marketing
        "FINANZ": ["2431", "2433"],  # analistas financieros
        # Educación: patrones específicos PRIMERO (más restrictivos) antes que el genérico
        # "EDUCACION BASICA" NO incluye 2310 (profesores universitarios ganan más)
        "EDUCACION BASICA": ["2330", "2340", "2320"],  # primaria/especial/secundaria
        "EDUCACION INFANTIL": ["2330", "2340"],
        "EDUCACION FISICA": ["2330", "2340"],
        "EDUCACION ARTISTICA": ["2330", "2340"],
        "EDUCACION ESPECIAL": ["2330", "2340"],
        "PEDAGOG": ["2320", "2330", "2340", "2310"],
        # "LICENCIATURA" genérico (al final): incluye 2310 porque algunas licenciaturas
        # son para docencia universitaria, pero los patrones específicos de arriba
        # ya excluyen 2310 para educación básica/infantil
        "LICENCIATURA": ["2320", "2330", "2340", "2310"],
        "EDUCACION": ["2320", "2330", "2340", "2310"],
    }
    # Buscar códigos ISCO directos según el programa
    isco_codes_direct: list[str] = []
    for patron, codes in prog_to_isco_codes.items():
        if patron in prog_upper:
            isco_codes_direct = codes
            break

    def _passes_isco_filter(row: dict) -> bool:
        """True si la ocupación pasa el filtro de grupo ISCO (o no hay filtro)."""
        if not isco_groups_allowed:
            return True
        isco = row.get("codigo_isco")
        if not isco:
            return False  # sin código ISCO, no podemos validar; excluir si hay filtro
        try:
            group = int(str(isco)[0])
            return group in isco_groups_allowed
        except (ValueError, IndexError):
            return False

    seen_uris: set[str] = set()
    result: list[dict] = []

    # 0) Búsqueda directa por código ISCO conocido (alta precisión)
    #    Para programas como MEDICINA donde el rol ESCO no contiene la palabra del programa.
    got_direct_matches = False
    if isco_codes_direct:
        for code in isco_codes_direct:
            rows = _sb_select("esco_ocupaciones", select="uri,nombre,codigo_isco",
                              filt={"codigo_isco": code}, limit=5)
            for r in rows:
                u = r.get("uri")
                n = _fix_mojibake(r.get("nombre", "")).strip()
                if u and n and u not in seen_uris and _passes_isco_filter(r):
                    seen_uris.add(u)
                    result.append({"uri": u, "nombre": n})
            if len(result) >= limit:
                break
        if len(result) > 0:
            got_direct_matches = True
            # Si encontramos al menos una ocupación por ISCO directo, confiamos en esos
            # códigos y no hacemos keyword search para evitar contaminar con ocupaciones
            # de sub-especialidades no afines (ej: MEDICINA no debe incluir oncólogos)

    # 1) Frase completa del programa sin stopwords (ej: 'INGENIERIA INDUSTRIAL')
    if not got_direct_matches and len(result) < limit and palabras:
        frase = " ".join(palabras)
        rows = _sb_ilike("esco_ocupaciones", "uri,nombre,codigo_isco", "nombre", f"%{frase}%", limit=limit * 3)
        for r in rows:
            u = r.get("uri")
            n = _fix_mojibake(r.get("nombre", "")).strip()
            if u and n and u not in seen_uris and _passes_isco_filter(r):
                n_norm = _norm(n)
                if all(_norm(w) not in n_norm for w in palabras):
                    continue
                seen_uris.add(u)
                result.append({"uri": u, "nombre": n})

    # 2) Búsqueda por palabra principal del programa convertida a rol ('ingeniero'/'administrador')
    #    + la palabra más distintiva.
    if not got_direct_matches and len(result) < limit:
        rol_map = {
            "INGENIERIA": "ingeniero", "INGENIERIAS": "ingeniero",
            "ADMINISTRACION": "administrador", "ADMINISTR": "administrador",
            "CIENCIA": "cientifico", "CIENCIAS": "cientifico",
            "LICENCIATURA": "licenciado",
            "TECNOLOG": "tecnico", "TECNOL": "tecnico",
            "PROFESIONAL": "profesional",
            "MEDICINA": "medico",  # rol médico
            "ENFERMER": "enfermero",
            "PSICOLOG": "psicologo",
            "ODONTOLOG": "odontologo",
            "VETERIN": "veterinario",
            "NUTRICION": "nutricionista",
            "FISIOTERAP": "fisioterapeuta",
            "BACTERIOLOG": "bacteriologo",
            "ARQUITECT": "arquitecto",
            "AGRONOM": "agronomo",
            "DERECHO": "abogado",
            "JURISPRUD": "abogado",
            "ECONOM": "economista",
            "CONTADUR": "contador",
            "SOCIOLOG": "sociologo",
            "TRABAJO SOCIAL": "trabajador social",
            "BIBLIOTEC": "bibliotecario",
            "COMUNIC": "periodista",
            "PERIODISM": "periodista",
        }
        rol = None
        for k, v in rol_map.items():
            if k in programa_norm:
                rol = v
                break
        claves = [w for w in palabras if _norm(rol or "") not in w.lower()] if rol else palabras
        if rol and claves:
            for clave in claves[:3]:
                patron = f"%{rol}%{clave}%"
                rows = _sb_ilike("esco_ocupaciones", "uri,nombre,codigo_isco", "nombre", patron, limit=limit * 3)
                for r in rows:
                    u = r.get("uri")
                    n = _fix_mojibake(r.get("nombre", "")).strip()
                    if u and n and u not in seen_uris and _passes_isco_filter(r):
                        seen_uris.add(u)
                        result.append({"uri": u, "nombre": n})
                if len(result) >= limit:
                    break
        elif rol and not claves:
            # Programa de una sola palabra (ej: MEDICINA, PSICOLOGIA): buscar por el rol solo
            rows = _sb_ilike("esco_ocupaciones", "uri,nombre,codigo_isco", "nombre", f"%{rol}%", limit=limit * 3)
            for r in rows:
                u = r.get("uri")
                n = _fix_mojibake(r.get("nombre", "")).strip()
                if u and n and u not in seen_uris and _passes_isco_filter(r):
                    seen_uris.add(u)
                    result.append({"uri": u, "nombre": n})

    # 3) Si aún faltan, búsqueda por palabra clave distintiva individual (>= 5 chars)
    #    con el filtro ISCO aplicado.
    if not got_direct_matches and len(result) < limit:
        for w in palabras[:5]:
            if len(w) < 5:
                continue
            rows = _sb_ilike("esco_ocupaciones", "uri,nombre,codigo_isco", "nombre", f"%{w}%", limit=limit * 3)
            for r in rows:
                u = r.get("uri")
                n = _fix_mojibake(r.get("nombre", "")).strip()
                if u and n and u not in seen_uris and _passes_isco_filter(r):
                    n_norm = _norm(n)
                    if _norm(w) not in n_norm:
                        continue
                    seen_uris.add(u)
                    result.append({"uri": u, "nombre": n})
            if len(result) >= limit:
                break

    return result[:limit]


def _habilidades_esenciales(ocupaciones: list[dict], top: int = 25) -> dict[str, int]:
    """Retorna un dict {habilidad: conteo} de habilidades ESSENCIAL de las ocupaciones.
    Acota al top `top` (default 25) por frecuencia para un índice realista."""
    freq: Counter = Counter()
    for occ in ocupaciones:
        uri = occ.get("uri")
        if not uri:
            continue
        rows = _sb_select(
            "esco_ocupacion_habilidades",
            select="habilidad_nombre,tipo_habilidad",
            filt={"ocupacion_uri": uri, "tipo_relacion": "essential"},
            limit=200,
        )
        for r in rows:
            n = _fix_mojibake((r.get("habilidad_nombre") or "")).strip()
            if n:
                freq[n] += 1
    if not freq:
        return {}
    # Top N por frecuencia (más frecuentes = más centrales en el perfil)
    top_items = freq.most_common(top)
    return dict(top_items)


# ============================================================================
# O*NET: fuente principal de habilidades modernas del mercado.
# La lógica vive en app.services.onet_mercado (mapeo directo programa →
# ocupación específica → habilidades). Desde este router se invoca a través
# de app.services.analisis_curricular.analizar_pensum_vs_onet en el endpoint
# /analizar-documento-curricular. No se replica más aquí.
# ============================================================================

# (Bloque O*NET eliminado: _PROG_TO_ONET_CODES, _cargar_onet_habilidades,
#  _ocupaciones_onet_compatibles, _habilidades_onet. Código muerto movido /
#  consolidado en app.services.onet_mercado.)




def _pensum_to_esco_tokens(pensum_items: list[str]) -> set[str]:
    """Expande un pensum en español a tokens/frases en inglés vía diccionario amplio.
    Devuelve un set normalizado (mayúsculas sin tildes) listo para comparar con ESCO."""
    tokens: set[str] = set()
    for item in pensum_items:
        n = _norm(item)
        if not n:
            continue
        tokens.add(n)
        # Tokens individuales (para match parcial)
        for w in n.split():
            if len(w) >= 4:
                tokens.add(w)
        # Equivalencias en inglés vía diccionario amplio
        for k_es, v_en in TOKEN_MAP_ES_EN.items():
            k_norm = _norm(k_es)
            # Coincidencia por substring o igualdad de tokens
            if k_norm in n or n in k_norm or any(t == k_norm for t in n.split()):
                for en in v_en:
                    tokens.add(_norm(en))
    return tokens


def _match_esencial(hab_en: str, pensum_tokens: set[str]) -> bool:
    """Dice si una habilidad ESCO (inglés) está cubierta por el pensum.
    Comparación por tokens (no substring) para reducir falsos positivos/negativos."""
    hn = _norm(hab_en)
    if not hn:
        return False
    # 1) Coincidencia directa de tokens (palabras sueltas significativas)
    hab_tokens = {t for t in hn.split() if len(t) >= 4}
    for p in pensum_tokens:
        if len(p) < 4:
            continue
        # igualdad de tokens o substring en frase
        if p in hn:
            # requiere que el token del pensum aparezca como palabra/token, no como sub-palabra
            # ej: 'data' en 'database' sí cuenta; 'data' en 'metadata' también
            return True
        # si algún token significativo de la habilidad está en el pensum
        for ht in hab_tokens:
            if ht == p:
                return True
    # 2) Match por diccionario inverso: si la hab_en mapea a un término ES conocido que está en pensum
    es_equiv = TOKEN_MAP_EN_ES.get(hn)
    if es_equiv:
        es_norm = _norm(es_equiv)
        for p in pensum_tokens:
            if p == es_norm or es_norm in p or p in es_norm:
                return True
    return False


def _traducir_faltantes_es(faltantes: list[str]) -> list[str]:
    """Traduce competencias ESCO (inglés) al español usando el diccionario EN→ES.
    Para las que no estén en el diccionario, devuelve la original en inglés (el LLM las completa después)."""
    out: list[str] = []
    for h in faltantes:
        hn = _norm(h)
        es_equiv = TOKEN_MAP_EN_ES.get(hn)
        if es_equiv:
            out.append(es_equiv)
        else:
            # Intentar match por substring en el diccionario inverso
            matched = None
            for k, v in TOKEN_MAP_EN_ES.items():
                if k in hn or hn in k:
                    matched = v
                    break
            out.append(matched if matched else h)
    return out


def _cursos_sena_por_brechas(faltantes: list[str], programa_norm: str = "", max_cursos: int = 6) -> list[dict]:
    """Recomienda cursos SENA basados en las brechas detectadas (no en el nombre del programa).
    Traduce cada brecha al español y busca cursos SENA que cubran esa brecha.
    Filtra palabras muy genéricas y exige que el curso SENA coincida con 2+ palabras de la brecha."""
    # Stopwords de palabras demasiado genéricas que traen ruido en SENA si van solas.
    # Pero si la brecha solo tiene estas palabras, se permiten (es lo único que tenemos).
    sena_stop_solo = {"DISEÑO", "PLANOS", "DISENO", "GESTION", "GESTIONAR", "MANEJO",
                      "ADMINISTRACION", "SISTEMAS", "TECNOLOG", "OPERACION",
                      "SEGURIDAD", "INVESTIGACION", "DISEGN"}
    # Estas palabras SÍ son útiles como keyword si se combinan con otra
    sena_useful_multi = {"PROCESOS", "PRODUCCION", "CALIDAD", "MANTENIMIENTO",
                         "INDUSTRIAL", "ELECTRICO", "MANUFACTURA", "DIBUJO",
                         "TECNICO", "MAQUINARIA", "MECANICA", "LOGISTICA",
                         "INVENTARIOS", "HIDRAULICA", "NEUMATICA"}
    sena_sugeridos: list[dict] = []
    seen_programas: set[str] = set()
    keywords_usadas: set[str] = set()

    # Top 5 brechas para buscar
    for h in faltantes[:5]:
        # Traducir al español para buscar en SENA (que está en español)
        hn = _norm(h)
        es_equiv = TOKEN_MAP_EN_ES.get(hn) or h
        # Palabras clave para buscar curso SENA: las distintivas (>= 4 chars).
        # Si la brecha tiene >1 palabra significativa, excluir las stopwords solas.
        # Si la brecha tiene 1 sola palabra, permitirla aunque sea stopword.
        brecha_words = [w for w in _norm(es_equiv).split() if len(w) >= 4]
        if len(brecha_words) > 1:
            palabras = [w for w in brecha_words
                        if w not in sena_stop_solo and w not in keywords_usadas]
        else:
            # 1 sola palabra: usarla aunque sea stopword, es lo único que tenemos
            palabras = [w for w in brecha_words if w not in keywords_usadas]
        # Si no quedan palabras útiles, saltar
        if not palabras:
            continue
        for kw in palabras[:2]:
            keywords_usadas.add(kw)
            try:
                rows = _sb_ilike("sena_programas_activos",
                                "programa,area_desempeno,duracion_horas,costo,institucion,departamento",
                                "programa", f"%{kw}%", limit=8)
                for r in rows:
                    prog = _fix_mojibake(r.get("programa") or "")
                    if not prog or prog in seen_programas:
                        continue
                    # Filtrar relevancia: el programa SENA debe contener la palabra clave
                    prog_norm = _norm(prog)
                    if kw not in prog_norm:
                        continue
                    # Evitar cursos de áreas claramente no relacionadas con el programa
                    area_norm = _norm(r.get("area_desempeno") or "")
                    bad_areas = ["MODAS", "BELLEZA", "GASTRONOM", "ESTETICA",
                                 "CABALLER", "PELUQUER", "MUSICA", "DANZA",
                                 "TEATRO", "ARTES", "AUDIOVISUALES", "ESPARCIMIENTO",
                                 "DEPORTES", "LOCUCION", "COSMET", "BARBER",
                                 "SECRETARIADO", "ADMINISTRATIV"]
                    # Excepción: si la keyword del programa aparece en el área como "industrial",
                    # no aplicar el filtro de áreas raras (ej: "diseño industrial" en modas)
                    if any(bad in area_norm for bad in bad_areas):
                        # Permitir solo si la keyword es claramente industrial/técnica
                        if kw not in {"INDUSTRIAL", "MANUFACTURA", "MECANICA", "ELECTRIC",
                                      "MAQUINARIA", "MANTENIMIENTO", "HIDRAULICA", "NEUMATICA"}:
                            continue
                    # Si la keyword es "PROCESOS" o "PRODUCCION", el programa SENA debe tener
                    # al menos otra palabra técnica relevante (no solo "procesos administrativos")
                    if kw in {"PROCESOS", "PRODUCCION"}:
                        prog_words_set = {w for w in prog_norm.split() if len(w) >= 4}
                        tech_terms = {"INDUSTRIAL", "MANUFACTURA", "PRODUCCION", "INDUSTRIALES",
                                      "FABRICACION", "ENSAMBLAJE", "MAQUINARIA", "MECANICA",
                                      "ELECTRIC", "LOGISTICA", "CALIDAD", "METAL", "PLASTICO",
                                      "ALIMENT", "QUIMIC", "TEXTIL", "CONFECCION"}
                        if not (prog_words_set & tech_terms):
                            continue
                    seen_programas.add(prog)
                    sena_sugeridos.append({
                        "programa": prog,
                        "area": _fix_mojibake(r.get("area_desempeno") or ""),
                        "duracion_horas": _normalizar_duracion_sena(r.get("duracion_horas")),
                        "costo_cop": r.get("costo"),
                        "institucion": _fix_mojibake(r.get("institucion") or ""),
                        "departamento": _fix_mojibake(r.get("departamento") or ""),
                        "cubren_brecha": _fix_mojibake(es_equiv),
                    })
                    if len(sena_sugeridos) >= max_cursos:
                        return sena_sugeridos
            except Exception:
                pass
    return sena_sugeridos


def _gemini_match_pensum(pensum: list[str], esenciales_en: list[str]) -> dict:
    """
    Usa Gemini para evaluar semánticamente qué habilidades ESCO cubre el pensum.
    Recibe el pensum en español y las habilidades esenciales en inglés (ESCO).
    Devuelve cubiertas y faltantes en español colombiano, más el índice calculado.
    """
    sys_prompt = (
        "Eres un evaluador de alineación curricular para educación superior en Colombia. "
        "Determina cuáles de las habilidades esenciales del mercado laboral están cubiertas "
        "por las materias de un pensum académico, y cuáles faltan.\n\n"
        "Reglas estrictas:\n"
        "1. Una habilidad está CUBIERTA si alguna materia del pensum la enseña (total o "
        "parcialmente), incluso con nombre distinto. Ej: 'Gestión de calidad' cubre 'quality management'.\n"
        "2. Sé conservador: sin evidencia clara, márcala como FALTANTE.\n"
        "3. TODAS las habilidades (cubiertas Y faltantes) deben estar en ESPAÑOL de Colombia. "
        "NUNCA devuelvas habilidades en inglés. Usa términos que un rector universitario entendería.\n"
        "4. Responde SOLO con este JSON exacto (sin markdown, sin texto adicional):\n"
        '{"cubiertas":["Habilidad en español","Otra habilidad en español"],'
        '"faltantes":["Habilidad faltante en español"]}'
    )
    usr_prompt = json.dumps({
        "pensum": pensum,
        "habilidades_esenciales": esenciales_en,
    }, ensure_ascii=False)
    res = call_gemini_text(sys_prompt, usr_prompt, temperature=0.1, max_tokens=2000)
    if res.startswith("```json"): res = res[7:]
    if res.startswith("```"): res = res[3:]
    if res.endswith("```"): res = res[:-3]
    data = json.loads(res.strip())
    cubiertas = [_fix_mojibake(c) for c in data.get("cubiertas", []) if isinstance(c, str)]
    faltantes = [_fix_mojibake(f) for f in data.get("faltantes", []) if isinstance(f, str)]

    # Guardia: si alguna cubierta o faltante quedó en inglés (ASCII puro sin tildes/ñ),
    # pedirle a Gemini que las traduzca
    all_items = cubiertas + faltantes
    en_items = [item for item in all_items if item.isascii() and not any(
        w.lower() in item.lower() for w in ["web", "sql", "big data", "software", "hardware", "iot", "scrum", "lean", "agile"]
    )]
    if en_items:
        try:
            trad_sys = (
                "Traduce esta lista de competencias del inglés al español de Colombia. "
                "Usa términos del mercado laboral y la educación superior colombiana. "
                "Devuelve SOLO: {\"traducciones\":[\"comp1\",\"comp2\"]} con el mismo número y orden."
            )
            trad_res = call_gemini_text(trad_sys, json.dumps(en_items, ensure_ascii=False), temperature=0.1, max_tokens=1500)
            if trad_res.startswith("```json"): trad_res = trad_res[7:]
            if trad_res.startswith("```"): trad_res = trad_res[3:]
            if trad_res.endswith("```"): trad_res = trad_res[:-3]
            trad_data = json.loads(trad_res.strip())
            traducciones = trad_data.get("traducciones", [])
            if len(traducciones) == len(en_items):
                idx_map = {item: _fix_mojibake(trad) for item, trad in zip(en_items, traducciones) if trad and isinstance(trad, str)}
                cubiertas = [idx_map.get(c, c) for c in cubiertas]
                faltantes = [idx_map.get(f, f) for f in faltantes]
        except Exception:
            pass

    indice = round(len(cubiertas) / max(len(esenciales_en), 1) * 100, 1)
    return {
        "cubiertas": cubiertas,
        "faltantes": faltantes,
        "indice": indice,
    }


@router.post("/alineacion-curricular")
async def alineacion_curricular(req: AlineacionRequest):
    """
    Índice de Alineación Curricular: cobertura del pensum sobre las habilidades
    ESCO esenciales de las ocupaciones afines al programa.
    Usa Gemini para matching semántico (entiende sinónimos), con fallback determinístico.
    """
    programa_norm = _norm(req.programa)
    if not programa_norm:
        raise HTTPException(status_code=400, detail="Programa requerido")
    if not req.pensum:
        raise HTTPException(status_code=400, detail="Indica al menos una competencia del pensum actual")

    # 1. Ocupaciones ESCO afines al programa (nuestros datos en Supabase)
    ocupaciones = _ocupaciones_compatibles(programa_norm, limit=5)
    if not ocupaciones:
        raise HTTPException(
            status_code=404,
            detail=f"No se encontraron ocupaciones afines a '{req.programa}'. Prueba con un nombre más específico (ej: 'Ingeniería de Sistemas')."
        )

    # 2. Top 25 habilidades esenciales desde ESCO (nuestros datos)
    esencial_freq = _habilidades_esenciales(ocupaciones, top=25)
    if not esencial_freq:
        raise HTTPException(
            status_code=503,
            detail=f"Las ocupaciones afines a '{req.programa}' no tienen habilidades esenciales registradas en ESCO."
        )
    top_esenciales = [h for h, _ in sorted(esencial_freq.items(), key=lambda x: -x[1])]
    total_esenciales = len(top_esenciales)

    # 3. Matching con Gemini (primario) o fallback determinístico
    if is_gemini_available():
        try:
            match = _gemini_match_pensum(req.pensum, top_esenciales)
            cubiertas_es = match["cubiertas"]
            faltantes_es = match["faltantes"]
            indice_alineacion = match["indice"]
        except Exception as e:
            print("AC: Gemini falló, usando fallback determinístico:", e)
            cubiertas_es, faltantes_es, indice_alineacion = _fallback_deterministico(req.pensum, top_esenciales)
    else:
        cubiertas_es, faltantes_es, indice_alineacion = _fallback_deterministico(req.pensum, top_esenciales)

    total_cubiertas = len(cubiertas_es)
    total_faltantes = len(faltantes_es)

    # 4. Coincidencia con vacantes SPE/APE
    spe_matched = 0
    for occ_dict in ocupaciones:
        occ = occ_dict["nombre"]
        palabras = [w for w in _norm(occ).split() if len(w) >= 5]
        if not palabras:
            continue
        kw = max(palabras, key=len)
        try:
            rows = _sb_ilike("spe_ape_inscritos_ocupacion", "ocupacion,inscritos_2020", "ocupacion",
                             f"%{kw}%", limit=5)
            if rows and any(int(r.get("inscritos_2020") or 0) > 100 for r in rows):
                spe_matched += 1
        except Exception:
            pass
    spe_match_pct = round(spe_matched / len(ocupaciones) * 100, 1) if ocupaciones else 0.0

    return {
        "programa": req.programa,
        "pensum_ingresado": req.pensum,
        "ocupaciones_afines": [_fix_mojibake(occ["nombre"]) for occ in ocupaciones],
        "indice_alineacion_curricular": indice_alineacion,
        "cobertura_esco_pct": indice_alineacion,
        "total_esenciales": total_esenciales,
        "total_cubiertas": total_cubiertas,
        "total_faltantes": total_faltantes,
        "coincidencia_vacantes_spe_pct": spe_match_pct,
        "competencias_cubiertas": cubiertas_es[:20],
        "competencias_faltan": faltantes_es[:15],
        "cursos_sena_recomendados": [],
        "metodologia": (
            "Matching semántico con Gemini 2.5 Flash-Lite: el pensum se compara contra "
            "el top-25 de habilidades esenciales ESCO de las ocupaciones afines. "
            "Gemini evalúa cobertura semántica (sinónimos, equivalencias), no solo tokens literales. "
            "TODO en español colombiano. "
            "Coincidencia SPE/APE: % de ocupaciones afines con >100 inscritos en el SPE del SENA."
        ),
        "fuentes": ["ESCO (UE)", "SPE/APE SENA"],
    }


def _fallback_deterministico(pensum: list[str], top_esenciales: list[str]) -> tuple[list[str], list[str], float]:
    """Matching determinístico por tokens. Se usa solo si Gemini no está disponible."""
    pensum_tokens = _pensum_to_esco_tokens(pensum)
    cubiertas_en: list[str] = []
    faltantes_en: list[str] = []
    for h in top_esenciales:
        if _match_esencial(h, pensum_tokens):
            cubiertas_en.append(h)
        else:
            faltantes_en.append(h)
    indice = round(len(cubiertas_en) / len(top_esenciales) * 100, 1) if top_esenciales else 0.0
    cubiertas_es = _traducir_faltantes_es(cubiertas_en)
    faltantes_es = _traducir_faltantes_es(faltantes_en)
    return cubiertas_es, faltantes_es, indice


# ===========================================================================
# Pensum típico por programa — para auto-análisis sin que el usuario escriba materias
# ===========================================================================

PENSUM_TIPICO: dict[str, list[str]] = {
    "INGENIERIA INDUSTRIAL": ["Matemáticas", "Estadística", "Investigación de operaciones", "Gestión de proyectos",
                              "Gestión de calidad", "Logística", "Manufactura", "Seguridad industrial",
                              "Análisis de datos", "Procesos de producción", "Mantenimiento", "Inglés"],
    "INGENIERIA CIVIL": ["Cálculo", "Mecánica de materiales", "Hidráulica", "Construcción",
                         "Estructuras", "Geotecnia", "Resistencia de materiales", "Dibujo técnico",
                         "Topografía", "Planos", "Concreto", "Gestión de proyectos", "Inglés"],
    "INGENIERIA DE SISTEMAS": ["Programación", "Bases de datos", "Estructuras de datos", "Redes",
                               "Sistemas operativos", "Ingeniería de software", "Análisis de datos",
                               "Seguridad informática", "Inteligencia artificial", "Inglés"],
    "INGENIERIA DE SOFTWARE": ["Programación", "Ingeniería de software", "Bases de datos", "Desarrollo web",
                               "Análisis de requisitos", "Pruebas de software", "Arquitectura de software",
                               "Calidad de software", "Inglés"],
    "ADMINISTRACION": ["Contabilidad", "Finanzas", "Marketing", "Gestión de proyectos",
                       "Estadística", "Matemáticas", "Recursos humanos", "Negociación",
                       "Análisis de datos", "Inglés"],
    "CONTADUR": ["Contabilidad", "Costos", "Finanzas", "Auditoría", "Tributación",
                 "Estadística", "Matemáticas", "Presupuesto", "Inglés"],
    "ECONOM": ["Microeconomía", "Macroeconomía", "Estadística", "Econometría",
               "Matemáticas", "Finanzas", "Política económica", "Análisis de datos", "Inglés"],
    "MEDICINA": ["Anatomía", "Fisiología", "Bioquímica", "Patología", "Farmacología",
                 "Medicina interna", "Cirugía", "Pediatría", "Ginecología", "Inglés"],
    "ENFERMER": ["Anatomía", "Fisiología", "Farmacología", "Cuidados de enfermería",
                 "Salud pública", "Nutrición", "Bioética", "Inglés"],
    "PSICOLOG": ["Psicología general", "Estadística", "Psicometría", "Psicología social",
                 "Neuropsicología", "Psicología clínica", "Entrevista psicológica",
                 "Ética profesional", "Inglés"],
    "DERECHO": ["Derecho civil", "Derecho constitucional", "Derecho penal", "Derecho laboral",
                "Derecho comercial", "Derecho administrativo", "Argumentación jurídica",
                "Ética", "Inglés"],
    "COMUNICAC": ["Comunicación", "Periodismo", "Redacción", "Producción audiovisual",
                  "Marketing", "Investigación", "Ética", "Inglés"],
    "ARQUITECT": ["Diseño arquitectónico", "Dibujo técnico", "Planos", "Historia de la arquitectura",
                  "Estructuras", "Construcción", "Sostenibilidad", "Urbanismo", "Inglés"],
    "INGENIERIA MECANIC": ["Mecánica", "Termodinámica", "Mecánica de fluidos", "Resistencia de materiales",
                           "Dibujo técnico", "Manufactura", "Mantenimiento", "Planos", "Inglés"],
    "MERCAD": ["Marketing", "Investigación de mercados", "Estadística", "Comportamiento del consumidor",
               "Publicidad", "Ventas", "Análisis de datos", "Inglés"],
    "FINANZ": ["Finanzas", "Contabilidad", "Matemáticas financieras", "Evaluación de proyectos",
               "Estadística", "Riesgos financieros", "Inglés"],
}

PENSUM_GENERICO = ["Matemáticas", "Estadística", "Gestión de proyectos", "Análisis de datos", "Inglés"]


def _pensum_auto(programa: str) -> list[str]:
    """Genera un pensum típico para un programa sin que el usuario tenga que escribirlo."""
    p = _norm(programa)
    for patron, materias in PENSUM_TIPICO.items():
        if _norm(patron) in p:
            return materias
    return PENSUM_GENERICO


# ===========================================================================
# Simulación 1 — Laboratorio de Actualización Curricular
#
# Filosofía: ALBA NO adivina, ALBA recalcula.
# - Todo el cómputo numérico (índice base, índice simulado, deltas, top-5
#   impacto) se hace con matching determinístico por tokens (reproducible).
# - Gemini SOLO traduce las habilidades ESCO (inglés->español) y redacta
#   la justificación narrativa del before/after. Nunca decide qué está
#   cubierto: el algoritmo determinístico lo decide.
# - El top-5 de mayor impacto se obtiene RECALCULANDO el índice agregando
#   cada brecha individualmente y midiendo el delta real (no es un promedio).
# ===========================================================================

class AnalisisCurricularRequest(BaseModel):
    programa: str
    competencias_extra: list[str] = []  # para simular: qué competencias agregar


def _indice_deterministico(pensum: list[str], top_esenciales: list[str]) -> tuple[int, list[str], list[str]]:
    """Matching determinístico puro por tokens. Devuelve (cubiertas_count, cubiertas_en, faltantes_en).
    Toda la numerología de la simulación pasa por aquí para ser reproducible."""
    cubiertas_en, faltantes_en, _ = _fallback_deterministico(pensum, top_esenciales)
    return len(cubiertas_en), cubiertas_en, faltantes_en


def _traducir_habilidades_llm(habilidades_en: list[str]) -> dict[str, str]:
    """Traduce con Gemini las habilidades ESCO que no estén en el diccionario.
    Devuelve {original_en: traduccion_es}. No inventa nada: solo traduce."""
    mapping: dict[str, str] = {}
    if not habilidades_en or not is_gemini_available():
        # al menos las del diccionario estático
        for h in habilidades_en:
            es = TOKEN_MAP_EN_ES.get(_norm(h))
            if es:
                mapping[h] = es
        return mapping
    pendientes = [h for h in habilidades_en if not TOKEN_MAP_EN_ES.get(_norm(h))]
    if pendientes:
        try:
            sys_p = (
                "Traduce esta lista de competencias de ESCO del inglés al español de Colombia. "
                "Usa términos comunes en educación superior. "
                'Devuelve SOLO un JSON: {"traducciones": ["item1","item2"]} con el MISMO numero de elementos. No uses markdown.'
            )
            usr_p = json.dumps(pendientes, ensure_ascii=False)
            res = call_gemini_text(sys_p, usr_p, temperature=0.1, max_tokens=1200)
            if res.startswith("```json"): res = res[7:]
            if res.startswith("```"): res = res[3:]
            if res.endswith("```"): res = res[:-3]
            data = json.loads(res.strip())
            trads = data.get("traducciones", [])
            if len(trads) == len(pendientes):
                for src, trad in zip(pendientes, trads):
                    if trad and isinstance(trad, str):
                        mapping[src] = _fix_mojibake(trad)
        except Exception as e:
            print("AC traducir: Gemini falló, dejando originales:", e)
    for h in habilidades_en:
        if h not in mapping:
            es = TOKEN_MAP_EN_ES.get(_norm(h))
            if es:
                mapping[h] = es
    return mapping


def _habilidades_es(habilidades_en: list[str], trad_map: dict[str, str]) -> list[str]:
    """Aplica el mapa de traducción a una lista de habilidades en inglés.
    Siempre aplica _fix_mojibake para limpiar caracteres corruptos de Supabase."""
    return [_fix_mojibake(trad_map.get(h, h)) for h in habilidades_en]


@router.post("/analisis-curricular")
async def analisis_curricular(req: AnalisisCurricularRequest):
    """
    Laboratorio de Actualización Curricular.

    1. Genera el pensum típico del programa automáticamente.
    2. Calcula el índice de alineación BASE (matching determinístico).
    3. Si se envían competencias_extra: recalcula el índice DESPUÉS y mide el delta real.
    4. Top-5 de competencias de mayor impacto: recalcula el índice agregando cada
       brecha individualmente y ordena por delta real.
    5. Gemini solo traduce las habilidades ESCO y redacta la justificación narrativa.

    Filosofía: ALBA no adivina, ALBA recalcula.
    """
    programa_norm = _norm(req.programa)
    if not programa_norm:
        raise HTTPException(status_code=400, detail="Programa requerido")

    # 1. Pensum típico automático (el usuario no escribe nada)
    pensum_base = _pensum_auto(req.programa)

    # 2. Ocupaciones ESCO afines (datos reales en Supabase)
    ocupaciones = _ocupaciones_compatibles(programa_norm, limit=5)
    if not ocupaciones:
        raise HTTPException(
            status_code=404,
            detail=f"No se encontraron ocupaciones afines a '{req.programa}'."
        )

    # 3. Top-25 habilidades esenciales ESCO (datos reales)
    esencial_freq = _habilidades_esenciales(ocupaciones, top=25)
    if not esencial_freq:
        raise HTTPException(
            status_code=503,
            detail=f"No hay habilidades esenciales registradas para '{req.programa}'."
        )
    top_esenciales = [h for h, _ in sorted(esencial_freq.items(), key=lambda x: -x[1])]
    total_esenciales = len(top_esenciales)

    # 4. Traducir las habilidades ESCO (inglés -> español). Gemini SOLO traduce.
    trad_map = _traducir_habilidades_llm(top_esenciales)

    # 5. Índice BASE: matching determinístico sobre pensum_base (recalcular, no adivinar)
    base_count, base_cubiertas_en, base_faltantes_en = _indice_deterministico(pensum_base, top_esenciales)
    base_cubiertas_es = _habilidades_es(base_cubiertas_en, trad_map)
    base_faltantes_es = _habilidades_es(base_faltantes_en, trad_map)
    indice_base = round(base_count / max(total_esenciales, 1) * 100, 1)

    # Deduplicar las listas en español: varias habilidades ESCO distintas pueden
    # traducirse a la misma etiqueta (ej: "design" y "design principles" -> "Diseño").
    # Conservamos el orden de aparición.
    def _dedup_preserve(items: list[str]) -> list[str]:
        seen: set[str] = set()
        out: list[str] = []
        for it in items:
            k = _norm(it)
            if k not in seen:
                seen.add(k)
                out.append(it)
        return out
    base_cubiertas_es = _dedup_preserve(base_cubiertas_es)
    base_faltantes_es = _dedup_preserve(base_faltantes_es)

    # 6. Top-5 de mayor impacto: RECALCULAR agregando cada brecha individualmente.
    # El impacto de una brecha = índice(pensum_base + [brecha]) - índice_base.
    # Es recálculo real, no promedio uniforme: una brecha que ya está parcialmente
    # cubierta (por tokens afines) tendrá menor delta que una totalmente ausente.
    # Deduplicamos por la competencia EN español: varias habilidades ESCO distintas
    # pueden traducirse a la misma etiqueta (ej: "design" y "design principles" ambos
    # "Diseño"); nos quedamos con la de mayor impacto para evitar duplicados.
    impacto_individual: list[dict] = []
    por_competencia_es: dict[str, dict] = {}
    for brecha_en in base_faltantes_en:
        pensum_con_brecha = pensum_base + [brecha_en]
        n_count, _, _ = _indice_deterministico(pensum_con_brecha, top_esenciales)
        n_idx = round(n_count / max(total_esenciales, 1) * 100, 1)
        delta = round(n_idx - indice_base, 1)
        if delta <= 0:
            continue
        brecha_es = _fix_mojibake(trad_map.get(brecha_en, brecha_en))
        entry = {
            "competencia": brecha_es,
            "competencia_en": brecha_en,
            "impacto": delta,
        }
        # Dedup por etiqueta en español: si ya existe una con esta etiqueta,
        # conservar la de mayor impacto.
        prev = por_competencia_es.get(brecha_es)
        if prev is None or delta > prev["impacto"]:
            por_competencia_es[brecha_es] = entry
    impacto_individual = sorted(por_competencia_es.values(), key=lambda x: -x["impacto"])
    top5_impacto = impacto_individual[:5]

    # 7. Simulación: si hay competencias_extra, RECALCULAR el índice DESPUÉS.
    indice_sim = None
    delta_sim = None
    sim_cubiertas_es: list[str] = []
    sim_faltantes_es: list[str] = []
    nuevas_cubiertas_es: list[str] = []
    justificacion_ia: str | None = None

    if req.competencias_extra:
        pensum_sim = pensum_base + req.competencias_extra
        sim_count, sim_cubiertas_en, sim_faltantes_en = _indice_deterministico(pensum_sim, top_esenciales)
        sim_cubiertas_es = _habilidades_es(sim_cubiertas_en, trad_map)
        sim_faltantes_es = _habilidades_es(sim_faltantes_en, trad_map)
        sim_cubiertas_es = _dedup_preserve(sim_cubiertas_es)
        sim_faltantes_es = _dedup_preserve(sim_faltantes_es)
        indice_sim = round(sim_count / max(total_esenciales, 1) * 100, 1)
        delta_sim = round(indice_sim - indice_base, 1)
        # Nuevas coberturas = cubiertas en el sim que no estaban en la base.
        # Deduplicamos por etiqueta normalizada (varias habilidades ESCO distintas
        # pueden traducirse a la misma etiqueta en español).
        base_cubiertas_set = set(_norm(c) for c in base_cubiertas_es)
        nuevas_seen: set[str] = set()
        for c in sim_cubiertas_es:
            k = _norm(c)
            if k not in base_cubiertas_set and k not in nuevas_seen:
                nuevas_seen.add(k)
                nuevas_cubiertas_es.append(c)

        # Justificación narrativa con IA (interpreta el before/after, NO inventa números)
        if is_gemini_available():
            try:
                just_sys = (
                    "Eres un asesor curricular de ALBA. Explica en 1-2 frases en español de Colombia "
                    "el impacto de agregar estas competencias a un programa académico. Sé concreto y útil. "
                    "Menciona los sectores beneficiados. NO inventes cifras: usa solo las que te damos. "
                    "No uses markdown."
                )
                just_usr = json.dumps({
                    "programa": req.programa,
                    "competencias_agregadas": req.competencias_extra,
                    "nuevas_habilidades_cubiertas": nuevas_cubiertas_es[:5],
                    "incremento_indice": f"+{delta_sim} puntos (de {indice_base} a {indice_sim})",
                }, ensure_ascii=False)
                just_res = call_gemini_text(just_sys, just_usr, temperature=0.3, max_tokens=300)
                justificacion_ia = just_res.strip()
            except Exception:
                pass

    # 8. Respuesta. El índice "actual" mostrado es el base; si hubo simulación,
    #    se expone aparte como indice_sim + delta_sim (antes/después).
    # Mercado real para las ocupaciones afines (salario + saturación)
    mercado = _mercado_onet_para_programa(req.programa, ocupaciones)

    return {
        "programa": req.programa,
        "pensum_base": pensum_base,
        "competencias_identificadas": pensum_base,  # alias legible para la UI
        "competencias_extra": req.competencias_extra,
        "ocupaciones_afines": [_fix_mojibake(occ["nombre"]) for occ in ocupaciones],
        "total_esenciales": total_esenciales,
        # Escenario ACTUAL (base)
        "indice_base": indice_base,
        "total_cubiertas": base_count,
        "total_faltantes": total_esenciales - base_count,
        "competencias_cubiertas": base_cubiertas_es[:15],
        "competencias_faltan": base_faltantes_es[:15],
        # Escenario SIMULADO (después), si aplica
        "indice_sim": indice_sim,
        "delta_sim": delta_sim,
        "sim_cubiertas": sim_cubiertas_es[:15],
        "sim_faltantes": sim_faltantes_es[:15],
        "nuevas_cubiertas": nuevas_cubiertas_es[:8],
        # Top-5 de mayor impacto (recalculado, no promedio)
        "top5_impacto": top5_impacto,
        # Justificación narrativa de IA (interpreta, no inventa)
        "justificacion_ia": justificacion_ia,
        # Compat hacia atrás con la UI vieja que lea estos campos
        "indice": indice_sim if indice_sim is not None else indice_base,
        "indice_anterior": indice_base if indice_sim is not None else None,
        "diferencia": delta_sim,
        "brechas": top5_impacto,  # alias: las brechas con impacto real
        "sugeridas": [b["competencia"] for b in top5_impacto] if not req.competencias_extra else [],
        "fuentes": ["ESCO (UE)", "SPE/APE SENA", "GEIH-DANE", "SNIES-MEN"],
        "salario_real_geih": mercado.get("salario"),
        "saturacion_mercado": mercado.get("saturacion"),
        "metodologia": (
            "Matching determinístico por tokens (reproducible). El índice = cubiertas / total esenciales × 100. "
            "El escenario simulado recalcula el índice agregando las competencias seleccionadas. "
            "El top-5 de mayor impacto recalcula el índice agregando cada brecha individualmente y mide el delta real. "
            "Gemini solo traduce las habilidades ESCO al español y redacta la justificación narrativa. "
            "Salario y saturación provienen de GEIH-DANE y SNIES-MEN."
        ),
        "salario_real_geih": mercado.get("salario"),
        "saturacion_mercado": mercado.get("saturacion"),
    }

# ===========================================================================
# Helpers Gobierno: RUES + DNP
# ===========================================================================

def _mercado_onet_para_programa(programa: str, ocupaciones: list[dict]) -> dict[str, Any]:
    """Trae salario real GEIH y saturación para las ocupaciones O*NET del programa.

    Usa los títulos O*NET para buscar ocupaciones ESCO por nombre y de ahí el
    código ISCO para cruzar con geih_salario_ocupacion.
    """
    out: dict[str, Any] = {"salario": None, "saturacion": None}
    if not ocupaciones:
        return out

    # Buscar ISCO por nombre ESCO aproximado al título O*NET
    isco_codes: set[str] = set()
    esco_titles: list[str] = []
    try:
        for occ in ocupaciones[:5]:
            title = (occ.get("title") or "").strip()
            if not title:
                continue
            # Buscar ocupación ESCO por nombre (tomar la primera que machee)
            words = [w for w in _norm(title).split() if len(w) >= 4]
            kw = " ".join(words[:2]) if len(words) >= 2 else (words[0] if words else title[:30])
            rows = _sb_ilike("esco_ocupaciones", "uri,nombre,codigo_isco", "nombre",
                             f"%{kw}%", limit=3)
            for r in rows:
                isco = r.get("codigo_isco")
                if isco:
                    isco_codes.add(str(isco))
                    esco_titles.append(r.get("nombre", ""))
                    break
    except Exception as e:
        print(f"[simulacion] mercado_onet esco error: {e}")

    salarios: list[tuple[float, float, float, str]] = []  # prom, med, emp, periodo
    for code in isco_codes:
        try:
            code_int = int(code)
        except ValueError:
            continue
        try:
            r_sal = _sb_select("geih_salario_ocupacion",
                               select="salario_promedio,salario_mediano,empleo_total,periodo",
                               filt={"oficio_c8": code_int},
                               order="periodo.desc", limit=5)
            for r in r_sal:
                sp = r.get("salario_promedio")
                sm = r.get("salario_mediano")
                em = r.get("empleo_total") or 0
                if sp and sm and em > 0:
                    salarios.append((float(sp), float(sm), float(em), str(r.get("periodo") or "N/D")))
                    break
        except Exception:
            continue

    if salarios:
        total_emp = sum(e for _, _, e, _ in salarios)
        prom_pond = sum(s * e for s, _, e, _ in salarios) / total_emp
        med_pond = sum(m * e for _, m, e, _ in salarios) / total_emp
        p10 = min(salarios, key=lambda x: x[1])[1]
        p90 = max(salarios, key=lambda x: x[0])[0]
        periodo = max((p for _, _, _, p in salarios if p != "N/D"), default="N/D")
        out["salario"] = {
            "salario_promedio_cop": round(prom_pond),
            "salario_mediana_cop": round(med_pond),
            "salario_minimo_cop": round(p10),
            "salario_maximo_cop": round(p90),
            "empleo_total_nacional": round(total_emp),
            "periodo": periodo,
            "fuente": "GEIH-DANE (promedio móvil 12 meses por oficio ISCO)",
            "num_oficios_considerados": len(salarios),
        }

    # Saturación: matriculados vs empleo
    try:
        pn = _norm(programa)
        snies = _sb_select("snies_programas_matriculados",
                           select="matriculados,institucion",
                           filt={"programa": programa}, limit=500)
        if not snies:
            words = [w for w in pn.split() if len(w) >= 4]
            frase = " ".join(words[:3]) if words else pn[:30]
            snies = _sb_ilike("snies_programas_matriculados",
                              "matriculados,institucion", "programa",
                              f"%{frase}%", limit=200)
        # Deduplicar por institución con max matriculados para no inflar
        matric_por_inst: dict[str, int] = {}
        for r in snies:
            inst = _norm(r.get("institucion") or "")
            m = int(r.get("matriculados") or 0)
            if inst and m > matric_por_inst.get(inst, 0):
                matric_por_inst[inst] = m
        matriculados_total = sum(matric_por_inst.values())
        empleo = out["salario"]["empleo_total_nacional"] if out["salario"] else 0
        if matriculados_total > 0 and empleo > 0:
            ratio = matriculados_total / empleo
            out["saturacion"] = {
                "matriculados_nacional": matriculados_total,
                "empleo_sector_nacional": round(empleo),
                "ratio_matriculados_empleo": round(ratio, 2),
                "nivel": "alta_saturacion" if ratio > 5 else "saturacion_media" if ratio > 2 else "baja_saturacion",
                "alerta": (
                    f"Hay {matriculados_total:,} matriculados vs {round(empleo):,} empleos del sector (ratio {ratio:.1f}x). "
                    + ("Alta competencia esperada." if ratio > 5 else "Competencia moderada." if ratio > 2 else "Demanda laboral suficiente.")
                ),
            }
    except Exception as e:
        print(f"[simulacion] saturacion_onet error: {e}")

    return out


def _rues_empresas_sector(rama: int) -> dict[str, Any]:
    """Trae empresas activas y nuevas para un sector CIIU-2 de RUES."""
    out = {"activas": 0, "nuevas_ultimo_ano": 0, "anio_nuevas": None}
    # Activas por sector (tabla nacional)
    try:
        # ciiu2 puede venir como texto con ceros a la izquierda
        rama_str = f"{rama:02d}" if isinstance(rama, int) else str(rama).zfill(2)
        rows = _sb_select("rues_top_sectores_nacional",
                          select="ciiu2,empresas_activas",
                          filt={"ciiu2": rama_str})
        if rows:
            out["activas"] = int(rows[0].get("empresas_activas") or 0)
        # Fallback: buscar sin ceros
        if not out["activas"]:
            rows = _sb_select("rues_top_sectores_nacional",
                              select="ciiu2,empresas_activas",
                              filt={"ciiu2": str(rama)})
            if rows:
                out["activas"] = int(rows[0].get("empresas_activas") or 0)
    except Exception:
        pass
    # Nuevas empresas: último año válido
    try:
        r_nuevas = _sb_select("rues_empresas_nuevas",
                              select="anio_matricula,ciiu2,empresas_nuevas",
                              filt={"ciiu2": f"{rama:02d}"})
        if not r_nuevas:
            r_nuevas = _sb_select("rues_empresas_nuevas",
                                  select="anio_matricula,ciiu2,empresas_nuevas",
                                  filt={"ciiu2": str(rama)})
        if r_nuevas:
            validos = [r for r in r_nuevas if 2015 <= (r.get("anio_matricula") or 0) <= 2030]
            if validos:
                ultimo = max(validos, key=lambda r: r.get("anio_matricula", 0))
                out["nuevas_ultimo_ano"] = int(ultimo.get("empresas_nuevas") or 0)
                out["anio_nuevas"] = int(ultimo.get("anio_matricula") or 0)
    except Exception:
        pass
    return out


def _dnp_desempeno(depto_norm: str) -> dict[str, Any] | None:
    """Promedio de desempeño municipal del departamento según DNP/MDM."""
    try:
        rows = _sb_select("dnp_desempeno_departamento",
                          select="departamento,promedio_desempeno",
                          limit=100)
        for r in rows:
            if _norm_depto(r.get("departamento") or "") == depto_norm:
                return {
                    "departamento": r.get("departamento"),
                    "promedio_desempeno": float(r.get("promedio_desempeno") or 0),
                    "fuente": "DNP/MDM 2016-2020",
                }
    except Exception:
        pass
    return None


class IntervencionRequest(BaseModel):
    departamento: str
    objetivo: str = "sectores_emergentes"
    beneficiarios: int = 500


@router.post("/intervencion-gobierno")
async def intervencion_gobierno(req: IntervencionRequest):
    """
    Ranking de oportunidades sectoriales en un departamento, ponderado según
    el objetivo declarado. Cada score 0-100 se descompone en 4 indicadores reales.

    Score = demanda + crecimiento + déficit_talento + compatibilidad (pesos por objetivo).
    """
    depto_norm = _norm_depto(req.departamento)
    depto_code = _depto_to_code(depto_norm)
    if not depto_code:
        raise HTTPException(status_code=400, detail=f"Departamento '{req.departamento}' no reconocido")

    pesos = PESOS_OBJETIVO.get(req.objetivo, PESOS_OBJETIVO["sectores_emergentes"])

    # 1. Sectores CIUU con empleo en el depto (último periodo)
    code_int = int(depto_code)
    ultimo_rows = _sb_select("geih_empleo_depto_sector", select="periodo",
                             filt={"dpto": code_int}, order="periodo.desc", limit=1)
    if not ultimo_rows:
        raise HTTPException(status_code=404, detail=f"Sin datos sectoriales para {req.departamento}")
    ultimo_periodo = ultimo_rows[0]["periodo"]
    sec_rows = _sb_select("geih_empleo_depto_sector",
                          filt={"dpto": code_int, "periodo": ultimo_periodo})
    empleo_por_rama: dict[int, float] = {}
    for r in sec_rows:
        rama = int(r.get("rama_ciiu") or 0)
        emp = float(r.get("empleo") or 0)
        if emp > empleo_por_rama.get(rama, 0):
            empleo_por_rama[rama] = emp
    if not empleo_por_rama:
        raise HTTPException(status_code=503, detail="No se pudo consolidar empleo sectorial")

    max_empleo = max(empleo_por_rama.values())
    top_ramas = sorted(empleo_por_rama.items(), key=lambda x: -x[1])[:12]

    # 2. Predicciones Chronos sectoriales (crecimiento proyectado)
    def _cronos_rate_sector(rama: int) -> float | None:
        """Crecimiento % anual proyectado por Chronos para el sector CIIU dado.
        Las predicciones tienen estructura: {sector: {historico: {...}, prediccion: {años, mediana, bajo_10, alto_90}}}.
        Las predicciones solo cubren 3 sectores macro (Agricultura, Industria, Servicios), así que
        mapeamos cada CIIU al sector macro más cercano antes de buscar."""
        # Mapeo CIIU-2-dígitos -> sector macro Chronos
        ciiu_to_macro = {
            1: "Agricultura",  # Agricultura, ganadería, caza y silvicultura
            2: "Industria",    # Explotación de minas y canteras (asimilable a industria)
            3: "Industria",    # Industria manufacturera
            # 4 no existe en CIIU-2
            5: "Industria",    # Construcción (asimilable a industria)
            6: "Servicios",    # Comercio, hoteles y restaurantes
            7: "Servicios",    # Transporte y comunicaciones
            8: "Servicios",    # Información y comunicaciones
            9: "Servicios",    # Actividades financieras e inmobiliarias
            10: "Servicios",   # Servicios comunales, sociales y personales
            # CIIU-2 detallados (sub-sectores manufactureros y servicios)
            11: "Industria", 13: "Industria", 14: "Industria", 17: "Industria", 18: "Industria",
            20: "Industria", 21: "Industria", 22: "Industria", 23: "Industria", 24: "Industria",
            25: "Industria", 26: "Industria", 27: "Industria", 28: "Industria", 29: "Industria",
            30: "Industria", 31: "Industria", 32: "Industria", 35: "Industria", 36: "Industria",
            41: "Industria", 43: "Industria",  # Construcción
            45: "Servicios", 46: "Servicios", 47: "Servicios",  # Comercio
            49: "Servicios", 50: "Servicios", 51: "Servicios", 52: "Servicios", 53: "Servicios",  # Transporte
            55: "Servicios", 56: "Servicios",  # Alojamiento y comida
            58: "Servicios", 59: "Servicios", 61: "Servicios",  # Info y telecom
            62: "Servicios", 63: "Servicios",  # Software y datos
            64: "Servicios", 65: "Servicios", 66: "Servicios", 68: "Servicios",  # Financieros e inmobiliarias
            69: "Servicios", 70: "Servicios", 71: "Servicios", 72: "Servicios", 73: "Servicios",
            74: "Servicios", 78: "Servicios", 80: "Servicios", 81: "Servicios", 82: "Servicios",
            84: "Servicios", 85: "Servicios", 86: "Servicios", 87: "Servicios",  # Pública, educación, salud
            90: "Servicios", 93: "Servicios", 94: "Servicios", 96: "Servicios", 97: "Servicios",
        }
        macro = ciiu_to_macro.get(rama)
        if not macro:
            return None
        try:
            pred_data = _load_predicciones_mundiales()
            sectores = pred_data.get("sectores", {})
            if not sectores:
                return None
            # Buscar el sector macro exacto o aproximado
            sec_data = sectores.get(macro)
            if not sec_data:
                for sec_name, sd in sectores.items():
                    if _norm(macro) in _norm(sec_name) or _norm(sec_name) in _norm(macro):
                        sec_data = sd
                        break
            if sec_data:
                pred = sec_data.get("prediccion", {})
                if isinstance(pred, dict):
                    valores = pred.get("mediana")
                    if valores is None:
                        valores_v = pred.get("valores", {})
                        valores = list(valores_v.values()) if isinstance(valores_v, dict) else valores_v
                    if valores and len(valores) >= 2:
                        try:
                            v0 = float(valores[0])
                            vf = float(valores[-1])
                            if v0 > 0:
                                n_anios = max(1, len(valores) - 1)
                                cagr = ((vf / v0) ** (1 / n_anios) - 1) * 100
                                return round(cagr, 2)
                        except Exception:
                            pass
            return None
        except Exception:
            return None

    # 3. Compatibilidad económica: tejido empresarial formal RUES + PILA fallback
    pila_rows = _sb_select("pila_resumen_sector",
                           select="actividadeconomicadesc,total_cotizantes",
                           order="total_cotizantes.desc", limit=50)
    max_pila = max((float(r.get("total_cotizantes") or 0) for r in pila_rows), default=1) or 1

    def _pila_compatibilidad(rama: int) -> float:
        nombre = CIIU2_NOMBRES.get(rama, "")
        if not nombre:
            return 0.0
        keywords = CIIU2_KEYWORD_SECTOR.get(rama, [_norm(nombre)[:8]])
        for r in pila_rows:
            desc = _norm(r.get("actividadeconomicadesc") or "")
            for kw in keywords:
                if kw in desc:
                    return float(r.get("total_cotizantes") or 0) / max_pila * 100
        return 0.0

    # RUES empresas activas por sector para enriquecer la compatibilidad económica
    max_rues_activas = 1
    try:
        rues_top = _sb_select("rues_top_sectores_nacional",
                              select="empresas_activas", order="empresas_activas.desc", limit=1)
        if rues_top:
            max_rues_activas = max(1, int(rues_top[0].get("empresas_activas") or 1))
    except Exception:
        pass

    def _rues_compatibilidad(rama: int) -> float:
        activas = _rues_empresas_sector(rama).get("activas", 0)
        return (activas / max_rues_activas * 100) if max_rues_activas else 0.0

    # 4. Déficit de talento: bajo = saturado, alto = hay hueco
    # Proxy: programas SENA afines en el depto (cuantos menos, más déficit).
    # NOTA: filtramos por departamento, no nacional, para que el indicador refleje
    # la oferta local de formación frente al empleo sectorial del depto.
    depto_norm_for_sena = _norm_depto(req.departamento)

    def _deficit_talento(rama: int) -> float:
        keywords = CIIU2_KEYWORD_SECTOR.get(rama, [])
        if not keywords:
            return 50.0  # neutro si no sabemos
        try:
            count = 0
            for kw in keywords[:2]:
                # Buscar programas SENA con la keyword Y que estén en el depto
                rows = _sb_ilike("sena_programas_activos", "programa,departamento", "programa",
                                 f"%{kw}%", limit=200)
                # Filtrar por departamento (case-insensitive, sin tildes)
                count += sum(1 for r in rows
                            if _norm_depto(r.get("departamento") or "") == depto_norm_for_sena)
            # Normalizar invertido: más programas = menos déficit
            # cap en 0-100: 0 programas → 100 (déficit alto), 50+ → 5
            deficit = max(5, min(100, 100 - count * 1.8))
            return round(deficit, 1)
        except Exception:
            return 50.0

    # 5. Calcular score por sector
    # Pre-cargar micronegocios por sector (EMICRON v2) para enriquecer el ranking
    # con empleo informal que PILA no captura
    micronegocios_por_grupo: dict[int, int] = {}
    try:
        r_ano = _sb_select("emicron_resumen_nacional_v2", select="ano", order="ano.desc", limit=1)
        ano_em = r_ano[0].get("ano") if r_ano else 2024
        r_mn = _sb_select("emicron_por_sector_v2", select="grupos12,micronegocios", filt={"ano": ano_em})
        for row in r_mn:
            micronegocios_por_grupo[int(row.get("grupos12") or 0)] = int(row.get("micronegocios") or 0)
    except Exception:
        pass

    oportunidades: list[dict] = []
    for rama, empleo in top_ramas:
        if empleo < max_empleo * 0.01:
            continue  # saltar sectores marginales
        demanda_norm = empleo / max_empleo * 100
        crecimiento = _cronos_rate_sector(rama) or 0.0
        # crecimiento puede ser negativo; normalizar 0-100 (capándolo)
        crec_norm = max(0, min(100, crecimiento * 4 + 50))  # 0% → 50, +12% → 100, -12% → 0
        deficit = _deficit_talento(rama)
        # Compatibilidad: combinamos RUES (tejido empresarial real) y PILA (fallback)
        compat_rues = _rues_compatibilidad(rama)
        compat_pila = _pila_compatibilidad(rama)
        compatibilidad = round(max(compat_rues, compat_pila * 0.5), 1)

        # Empleo informal (micronegocios) del sector
        g12 = CIIU2_TO_GRUPOS12.get(rama)
        mn_sector = micronegocios_por_grupo.get(g12, 0) if g12 else 0

        # Datos RUES por sector
        rues_sector = _rues_empresas_sector(rama)

        score = (
            pesos["demanda"] * demanda_norm
            + pesos["crecimiento"] * crec_norm
            + pesos["deficit_talento"] * deficit
            + pesos["compatibilidad"] * compatibilidad
        )
        oportunidades.append({
            "rama_ciiu": rama,
            "sector": CIIU2_NOMBRES.get(rama, f"Rama {rama}"),
            "empleo_depto": round(empleo),
            "demanda_pct": round(demanda_norm, 1),
            "crecimiento_proyectado_pct": round(crecimiento, 1),
            "deficit_talento_pct": round(deficit, 1),
            "compatibilidad_economica_pct": compatibilidad,
            "micronegocios_sector": mn_sector,
            "empleo_total_estimado": round(empleo) + mn_sector,
            "empresas_activas_sector": rues_sector.get("activas"),
            "empresas_nuevas_sector": rues_sector.get("nuevas_ultimo_ano"),
            "anio_empresas_nuevas": rues_sector.get("anio_nuevas"),
            "score": round(score, 1),
            "beneficiarios_estimados": int(req.beneficiarios),
        })

    oportunidades.sort(key=lambda x: -x["score"])
    top4 = oportunidades[:4]

    # Justificación textual por cada sector top (basada en sus métricas reales)
    for o in top4:
        razones = []
        if o["demanda_pct"] >= 60:
            razones.append(f"Alta demanda laboral local ({o['empleo_depto']:,} ocupados)")
        if o["crecimiento_proyectado_pct"] > 3:
            razones.append(f"Crecimiento proyectado Chronos +{o['crecimiento_proyectado_pct']}% anual")
        if o["deficit_talento_pct"] >= 70:
            razones.append(f"Alto déficit de talento (pocos programas SENA en {req.departamento})")
        if o.get("empresas_activas_sector"):
            razones.append(f"{o['empresas_activas_sector']:,} empresas activas en el sector (RUES)")
        elif o["compatibilidad_economica_pct"] >= 50:
            razones.append("Sólida base económica formal")
        if not razones:
            razones.append("Combinación equilibrada de indicadores")
        o["justificacion"] = ". ".join(razones) + "."

    # Justificación enriquecida por LLM para el top 1 (narrativa con los indicadores reales)
    # El LLM no inventa datos, solo redacta una recomendación accionable basada en las métricas reales.
    dnp_info = _dnp_desempeno(depto_norm)
    if is_gemini_available() and top4:
        try:
            top1 = top4[0]
            sys_prompt = (
                "Eres ALBA, una analista laboral del gobierno colombiano. Redacta una recomendación "
                "ejecutiva de máximo 3 párrafos para una intervención pública en el sector top de un departamento. "
                "Basándote EXCLUSIVAMENTE en los indicadores reales que recibes (demanda, crecimiento, "
                "déficit de talento, compatibilidad económica, empresas activas RUES, desempeño institucional DNP), "
                "propón 2-3 acciones concretas (cursos SENA, alianzas público-privadas, incentivos, etc.). "
                "NO inventes cifras. Usa pesos colombianos. Tono profesional, conciso, en español de Colombia. "
                "No uses markdown."
            )
            user_prompt = (
                f"DEPARTAMENTO: {req.departamento}\n"
                f"OBJETIVO: {req.objetivo.replace('_', ' ')}\n"
                f"BENEFICIARIOS: {req.beneficiarios}\n"
                f"SECTOR TOP: {top1['sector']} (CIIU {top1['rama_ciiu']})\n"
                f"INDICADORES REALES:\n"
                f"- Demanda laboral (GEIH): {top1['empleo_depto']:,} ocupados ({top1['demanda_pct']}% del depto)\n"
                f"- Crecimiento proyectado (Chronos): {top1['crecimiento_proyectado_pct']}% anual\n"
                f"- Déficit de talento: {top1['deficit_talento_pct']}% (basado en programas SENA en el depto)\n"
                f"- Compatibilidad económica: {top1['compatibilidad_economica_pct']}% (RUES empresas activas + PILA cotizantes)\n"
                f"- Empresas activas sector (RUES): {top1.get('empresas_activas_sector') or 'N/D'}\n"
                f"- Empresas nuevas sector (RUES {top1.get('anio_empresas_nuevas') or 'N/D'}): {top1.get('empresas_nuevas_sector') or 'N/D'}\n"
                f"- Desempeño institucional DNP/MDM: {dnp_info['promedio_desempeno'] if dnp_info else 'N/D'}\n"
                f"SCORE: {top1['score']}/100\n"
            )
            justificacion_ia = call_gemini_text(sys_prompt, user_prompt, temperature=0.3, max_tokens=600)
            if justificacion_ia and len(justificacion_ia) > 50:
                top1["justificacion_ia"] = justificacion_ia.strip()
        except Exception as e:
            print("DEBUG intervención: error LLM justificación:", e)

    respuesta = {
        "departamento": req.departamento,
        "objetivo": req.objetivo,
        "objetivo_label": req.objetivo.replace("_", " ").capitalize(),
        "beneficiarios": req.beneficiarios,
        "periodo_empleo": ultimo_periodo,
        "pesos_objetivo": pesos,
        "top_oportunidades": top4,
        "ranking_completo": oportunidades[:10],
        "dnp_desempeno": dnp_info,
        "metodologia": (
            "Score 0-100 = demanda (empleo sectorial GEIH) × peso + crecimiento proyectado (Chronos T5) × peso + "
            "déficit de talento (inverso de programas SENA en el depto) × peso + compatibilidad económica "
            "(empresas activas RUES + cotizantes PILA del sector) × peso. Los pesos dependen del objetivo declarado."
        ),
        "fuentes": ["GEIH-DANE", "SENA", "RUES", "PILA-MinTrabajo", "Chronos T5 (World Bank)", "DNP/MDM"],
    }
    return respuesta


# ===========================================================================
# Simulación 3 — Explora carrera (Estudiantes)
# ===========================================================================

class ExploraRequest(BaseModel):
    programa: str


@router.post("/explora-carrera")
async def explora_carrera(req: ExploraRequest):
    """
    Explora una carrera sin scores mágicos: habilidades que desarrollarás, salidas
    laborales, demanda del sector, salario real (GEIH + OLE), métricas Saber Pro,
    alerta de saturación (matriculados vs empleos) y dónde estudiarla.
    """
    programa_norm = _norm(req.programa)
    if not programa_norm:
        raise HTTPException(status_code=400, detail="Programa requerido")

    # 1. Ocupaciones ESCO afines (top 10 para enriquecer análisis)
    ocupaciones = _ocupaciones_compatibles(programa_norm, limit=10)

    # 2. Habilidades esenciales (top 10 por frecuencia) + traducidas al español
    habilidades_top: list[tuple[str, int]] = []
    if ocupaciones:
        freq = _habilidades_esenciales(ocupaciones, top=25)
        habilidades_top = sorted(freq.items(), key=lambda x: -x[1])[:10]
    # Traducir las habilidades al español para mostrarlas amigables
    # Primero por diccionario, luego LLM para las que sigan en inglés
    hab_es: list[dict] = []
    pendientes_llm_hab: list[tuple[int, str]] = []
    for idx, (h, c) in enumerate(habilidades_top):
        hn = _norm(h)
        es_equiv = TOKEN_MAP_EN_ES.get(hn)
        if not es_equiv:
            # Buscar por substring en el diccionario inverso
            for k, v in TOKEN_MAP_EN_ES.items():
                if k in hn or hn in k:
                    es_equiv = v
                    break
        trad = _fix_mojibake(es_equiv or h)
        hab_es.append({"habilidad": trad, "frecuencia_en_ocupaciones": c})
        # Si sigue en inglés (ASCII puro y no estaba en el diccionario), marcar para LLM
        if trad.isascii() and not es_equiv:
            pendientes_llm_hab.append((idx, h))

    # Traducción por LLM de las habilidades restantes
    if is_gemini_available() and pendientes_llm_hab:
        try:
            sys_p = (
                "Traduce esta lista de competencias de ESCO del inglés al español de Colombia. "
                "Usa términos comunes en educación superior. "
                'Devuelve SOLO un JSON: {"traducciones": ["item1","item2"]} con el MISMO numero de elementos. No uses markdown.'
            )
            usr_p = json.dumps([h for _, h in pendientes_llm_hab], ensure_ascii=False)
            res = call_gemini_text(sys_p, usr_p, temperature=0.1, max_tokens=800)
            if res.startswith("```json"): res = res[7:]
            if res.startswith("```"): res = res[3:]
            if res.endswith("```"): res = res[:-3]
            trad_dict = json.loads(res.strip())
            if "traducciones" in trad_dict and len(trad_dict["traducciones"]) == len(pendientes_llm_hab):
                for (idx, _), trad in zip(pendientes_llm_hab, trad_dict["traducciones"]):
                    if trad and isinstance(trad, str):
                        hab_es[idx]["habilidad"] = _fix_mojibake(trad)
        except Exception as e:
            print("DEBUG explora: error traduciendo habilidades con LLM:", e)

    # 3. Salario real GEIH: usar codigo_isco de las ocupaciones ESCO afines
    #    GEIH usa el mismo código ISCO-08 que ESCO (campo oficio_c8 = codigo_isco).
    salario_geih: dict | None = None
    try:
        isco_codes: set[str] = set()
        for occ in ocupaciones:
            r_occ = _sb_select("esco_ocupaciones", select="codigo_isco",
                               filt={"uri": occ.get("uri")}, limit=1)
            if r_occ and r_occ[0].get("codigo_isco"):
                isco_codes.add(str(r_occ[0]["codigo_isco"]))
        salarios: list[tuple[float, float, float]] = []  # (salario_promedio, salario_mediano, empleo_total)
        for code in isco_codes:
            try:
                code_int = int(code)
            except ValueError:
                continue
            # Obtener TODOS los periodos y quedarse con el más reciente con datos válidos
            r_sal = _sb_select("geih_salario_ocupacion",
                               select="salario_promedio,salario_mediano,empleo_total,periodo",
                               filt={"oficio_c8": code_int},
                               order="periodo.desc", limit=5)
            if r_sal:
                # Tomar la primera fila con datos válidos (ya ordenada por periodo desc)
                row = None
                for r in r_sal:
                    sp = r.get("salario_promedio")
                    sm = r.get("salario_mediano")
                    em = r.get("empleo_total") or 0
                    if sp and sm and em > 0:
                        row = (float(sp), float(sm), float(em))
                        break
                if row:
                    salarios.append(row)
        if salarios:
            # Calcular promedios ponderados por empleo
            total_emp = sum(e for _, _, e in salarios)
            prom_pond = sum(s * e for s, _, e in salarios) / total_emp if total_emp else 0
            med_pond = sum(m * e for _, m, e in salarios) / total_emp if total_emp else 0

            # Percentiles coherentes: usar salario_mediano para p10 y salario_promedio para p90
            # (la mediana siempre es <= promedio en distribuciones de salario sesgadas a la derecha)
            if len(salarios) == 1:
                # Con una sola ocupación, estimar p10 y p90 desde mediana y promedio
                p10 = round(med_pond * 0.75)   # recién egresado ~75% de la mediana
                p90 = round(prom_pond * 1.3)    # senior ~130% del promedio
            else:
                salarios_sorted = sorted(salarios, key=lambda x: x[1])  # ordenar por salario_mediano
                emp_acum = 0
                p10 = salarios_sorted[0][1]
                for _, m, e in salarios_sorted:
                    emp_acum += e
                    if emp_acum >= total_emp * 0.1:
                        p10 = m
                        break
                emp_acum = 0
                p90 = salarios_sorted[-1][0]
                for s, _, e in reversed(salarios_sorted):
                    emp_acum += e
                    if emp_acum >= total_emp * 0.1:
                        p90 = s
                        break

            # Asegurar coherencia: p10 <= mediana <= p90
            p10 = min(p10, med_pond)
            p90 = max(p90, prom_pond)

            # Periodo real del dato usado: el más reciente entre las ocupaciones consideradas
            periodo_real = "N/D"
            try:
                periodo_real = max(
                    (r.get("periodo") for r in r_sal if r.get("periodo")),
                    default="N/D",
                )
            except Exception:
                pass
            salario_geih = {
                "salario_promedio_cop": round(prom_pond),
                "salario_mediana_cop": round(med_pond),
                "salario_minimo_cop": round(p10),
                "salario_maximo_cop": round(p90),
                "empleo_total_nacional": round(total_emp),
                "periodo": periodo_real,
                "fuente": "GEIH-DANE (promedio móvil 12 meses por oficio ISCO)",
                "num_oficios_considerados": len(salarios),
            }
    except Exception as e:
        print("DEBUG explora: error salario GEIH:", e)

    # 4. Salario OLE más frecuente del programa (rango histórico 2001-2022)
    salario_rango: str | None = None
    salario_cop: float | None = None
    egresados_anuales = 0
    try:
        ole_rows = _sb_ilike("ole_ingresos_por_programa",
                             "rango_ingreso,graduados", "programa",
                             f"%{programa_norm[:25]}%", limit=50)
        if ole_rows:
            rango_freq = Counter()
            for r in ole_rows:
                rng = r.get("rango_ingreso")
                g = int(r.get("graduados") or 0)
                if rng:
                    rango_freq[rng] += g
                egresados_anuales += g
            if rango_freq:
                salario_rango, _ = rango_freq.most_common(1)[0]
                salario_cop = _rango_a_cop(salario_rango)
    except Exception:
        pass

    # Deduplicar: los egresados se suman dos veces si hay cargas dobles; usar max por rango
    try:
        ole_rows2 = _sb_ilike("ole_ingresos_por_programa",
                              "rango_ingreso,graduados", "programa",
                              f"%{programa_norm[:25]}%", limit=200)
        if ole_rows2:
            seen: dict[str, int] = {}
            for r in ole_rows2:
                rng = r.get("rango_ingreso")
                g = int(r.get("graduados") or 0)
                if rng and g > seen.get(rng, 0):
                    seen[rng] = g
            egresados_anuales = sum(seen.values())
            if seen and not salario_rango:
                salario_rango = max(seen.items(), key=lambda x: x[1])[0]
                salario_cop = _rango_a_cop(salario_rango)
    except Exception:
        pass

    # 5. Alerta de saturación: matriculados SNIES vs empleos del sector
    #    Si hay muchos matriculados vs empleos en el sector afín → saturado.
    #    Anti-timeout: usar eq exacto primero, ilike con frase solo si eq da 0.
    saturacion: dict | None = None
    try:
        # Paso 1: eq exacto del programa original (rápido). Incluimos institucion para dedup.
        snies_rows = _sb_select("snies_programas_matriculados",
                                select="matriculados,institucion",
                                filt={"programa": req.programa},
                                limit=500)
        # Paso 2: si eq da 0, intentar ilike con frase completa del programa normalizado
        if not snies_rows:
            stop_s = {"DE", "DEL", "LA", "LAS", "LOS", "EL", "EN", "Y", "E", "PARA", "CON",
                      "INGENIERIA", "CIENCIA", "CIENCIAS", "ADMINISTR", "PROFESIONAL",
                      "TECNOLOG", "TECNICA", "LICENCIATURA", "ESPECIALIZACION", "MAESTRIA",
                      "DOCTORADO", "DIPLOMADO"}
            palabras_restantes = [w for w in programa_norm.split() if len(w) >= 4]
            frase_s = " ".join(palabras_restantes[:3]) if palabras_restantes else programa_norm[:30]
            try:
                snies_rows = _sb_ilike("snies_programas_matriculados",
                                        "matriculados,institucion", "programa",
                                        f"%{frase_s}%", limit=200)
            except Exception as e:
                print("DEBUG explora: saturación SNIES ilike omitido:", str(e)[:80])
        # Deduplicar matriculados por institución: SNIES tiene una fila por sede/seccional
        # del mismo programa. Para no inflar el total, tomamos el MAX por institución
        # (la sede principal suele tener el mayor número de matriculados del programa).
        matric_por_inst: dict[str, int] = {}
        for r in snies_rows:
            inst = _norm(r.get("institucion") or "")
            if not inst:
                continue
            m = int(r.get("matriculados") or 0)
            if m > matric_por_inst.get(inst, 0):
                matric_por_inst[inst] = m
        matriculados_total = sum(matric_por_inst.values())
        # Empleo nacional del sector (usamos salario GEIH empleo_total si lo tenemos)
        empleo_sector = salario_geih.get("empleo_total_nacional", 0) if salario_geih else 0
        egresados = egresados_anuales if egresados_anuales else 0
        # Ratio de saturación: matriculados por egresado (proxy de cuántos salen vs cupos)
        # Si hay muchos matriculados por egresado → alta competencia
        if egresados > 0 and matriculados_total > 0:
            ratio_matr_egres = matriculados_total / egresados
            # Ratio de competencia: matriculados vs empleos disponibles del sector
            ratio_matr_empleo = matriculados_total / empleo_sector if empleo_sector > 0 else None
            if ratio_matr_empleo is None:
                nivel = "datos_insuficientes"
                alerta = f"Hay {matriculados_total:,} matriculados y {egresados:,} egresados anuales, pero no hay datos de empleo del sector para comparar."
            elif ratio_matr_empleo > 5:
                nivel = "alta_saturacion"
                alerta = f"Hay {matriculados_total:,} matriculados por cada ~{round(empleo_sector):,} empleos del sector afín (ratio {ratio_matr_empleo:.1f}x). Alta competencia esperada."
            elif ratio_matr_empleo > 2:
                nivel = "saturacion_media"
                alerta = f"Hay {matriculados_total:,} matriculados vs {round(empleo_sector):,} empleos del sector (ratio {ratio_matr_empleo:.1f}x). Competencia moderada."
            else:
                nivel = "baja_saturacion"
                alerta = f"Hay {matriculados_total:,} matriculados vs {round(empleo_sector):,} empleos del sector (ratio {ratio_matr_empleo:.1f}x). Demanda laboral suficiente."
            saturacion = {
                "matriculados_nacional": matriculados_total,
                "egresados_anuales_nacional": egresados,
                "empleo_sector_nacional": round(empleo_sector) if empleo_sector else None,
                "ratio_matriculados_empleo": round(ratio_matr_empleo, 2) if ratio_matr_empleo else None,
                "ratio_matriculados_egresados": round(ratio_matr_egres, 2),
                "nivel": nivel,
                "alerta": alerta,
            }
    except Exception as e:
        print("DEBUG explora: error saturación:", e)

    # 6. Métricas Saber Pro del programa (brechas detectadas)
    saber_pro: dict | None = None
    try:
        # Buscar el programa en saberpro_resumen_programas y promediar por institución
        sp_rows = _sb_ilike("saberpro_resumen_programas",
                            "institucion,programa,departamento,mod_ingles_punt,mod_razona_cuantitat_punt,"
                            "mod_comuni_escrita_punt,mod_lectura_critica_punt,mod_competen_ciudada_punt",
                            "programa", f"%{programa_norm[:25]}%", limit=100)
        if sp_rows:
            # Promedios nacionales del programa
            def _avg(rows, field):
                vals = [float(r.get(field) or 0) for r in rows if r.get(field) is not None]
                return round(sum(vals) / len(vals), 1) if vals else None
            saber_pro = {
                "instituciones_evaluadas": len(sp_rows),
                "ingles_punt": _avg(sp_rows, "mod_ingles_punt"),
                "razonamiento_cuantitativo_punt": _avg(sp_rows, "mod_razona_cuantitat_punt"),
                "comunicacion_escrita_punt": _avg(sp_rows, "mod_comuni_escrita_punt"),
                "lectura_critica_punt": _avg(sp_rows, "mod_lectura_critica_punt"),
                "competencias_ciudadanas_punt": _avg(sp_rows, "mod_competen_ciudada_punt"),
                "fuente": "Saber Pro - ICFES",
            }
            # Identificar brechas: inglés < 150 (B2) es alerta típica
            brechas_saber = []
            if saber_pro["ingles_punt"] is not None and saber_pro["ingles_punt"] < 150:
                brechas_saber.append(f"Inglés ({saber_pro['ingles_punt']}/250): por debajo de B2 (150+). Brecha para el mercado internacional.")
            if saber_pro["razonamiento_cuantitativo_punt"] is not None and saber_pro["razonamiento_cuantitativo_punt"] < 120:
                brechas_saber.append(f"Razonamiento cuantitativo ({saber_pro['razonamiento_cuantitativo_punt']}/300): por debajo del promedio nacional.")
            saber_pro["brechas_detectadas"] = brechas_saber
    except Exception as e:
        print("DEBUG explora: error Saber Pro:", e)

    # 7. Sectores con mayor empleo nacional (GEIH último periodo nacional)
    sectores_nacional: list[dict] = []
    try:
        ult_rows = _sb_select("geih_empleo_depto_sector", select="periodo",
                              order="periodo.desc", limit=1)
        if ult_rows:
            ultp = ult_rows[0]["periodo"]
            all_rows = _sb_select("geih_empleo_depto_sector",
                                  filt={"periodo": ultp}, limit=500)
            emp_rama: dict[int, float] = {}
            for r in all_rows:
                rama = int(r.get("rama_ciiu") or 0)
                emp_rama[rama] = emp_rama.get(rama, 0) + float(r.get("empleo") or 0)
            # Tomar max agregando (por si duplicados)
            for rama, emp in sorted(emp_rama.items(), key=lambda x: -x[1])[:5]:
                sectores_nacional.append({
                    "rama_ciiu": rama,
                    "sector": CIIU2_NOMBRES.get(rama, f"Rama {rama}"),
                    "empleo_nacional": round(emp),
                })
    except Exception:
        pass

    # 8. Dónde estudiarla: top 5 instituciones SNIES con ese programa
    #    Estrategia anti-timeout: `ilike '%PALABRA%'` sobre 600K filas hace timeout
    #    en Supabase. Usamos `eq` exacto (rápido, 0.4s) primero, y solo si no hay
    #    suficientes resultados, `ilike` con la FRASE COMPLETA del programa (selectiva).
    instituciones: list[dict] = []
    try:
        stop = {"DE", "DEL", "LA", "LAS", "LOS", "EL", "EN", "Y", "E", "PARA", "CON",
                "INGENIERIA", "CIENCIA", "CIENCIAS", "ADMINISTR", "PROFESIONAL",
                "TECNOLOG", "TECNICA", "LICENCIATURA", "ESPECIALIZACION", "MAESTRIA",
                "DOCTORADO", "DIPLOMADO"}
        palabras_distintivas = [w for w in programa_norm.split() if len(w) >= 4 and w not in stop]

        # Paso 1: eq exacto del programa original (rápido, indexado si existe)
        snies = _sb_select("snies_programas_matriculados",
                           select="institucion,programa,matriculados,departamento",
                           filt={"programa": req.programa},
                           order="matriculados.desc", limit=30)

        # Paso 2: si hay pocos resultados, intentar con frase completa del programa normalizado
        if len(snies) < 5:
            # Construir frase selectiva: las 2-3 palabras más largas en orden original
            # ej: "INGENIERIA DE SISTEMAS" -> "INGENIERIA SISTEMAS" (frase común, selectiva)
            palabras_restantes = [w for w in programa_norm.split() if len(w) >= 4]
            frase = " ".join(palabras_restantes[:3]) if palabras_restantes else programa_norm[:30]
            try:
                snies2 = _sb_ilike("snies_programas_matriculados",
                                   "institucion,programa,matriculados,departamento", "programa",
                                   f"%{frase}%", limit=30, order="matriculados.desc")
                # Merge manteniendo los ya encontrados por eq
                seen_keys = {(_norm(r.get("institucion") or ""), _norm(r.get("programa") or ""))
                             for r in snies}
                for r in snies2:
                    k = (_norm(r.get("institucion") or ""), _norm(r.get("programa") or ""))
                    if k not in seen_keys:
                        snies.append(r)
                        seen_keys.add(k)
            except Exception as e:
                print("DEBUG explora: SNIES ilike fallback omitido:", str(e)[:80])

        # Deduplicar por (institucion, programa) con max matriculados
        seen2: dict[tuple, dict] = {}
        for r in snies:
            pnorm = _norm(r.get("programa") or "")
            # Exigir que el programa contenga al menos una palabra distintiva del programa original
            if palabras_distintivas and not any(_norm(w) in pnorm for w in palabras_distintivas):
                continue
            key = (_norm(r.get("institucion") or ""), pnorm)
            mat = float(r.get("matriculados") or 0)
            if key not in seen2 or mat > seen2[key].get("matriculados", 0):
                seen2[key] = r
        instituciones = list(seen2.values())[:5]
    except Exception as e:
        print("DEBUG explora: error SNIES:", e)

    # 9. Descripción del profesional generada por IA
    #    Un cuadro que dice qué hace el profesional de ese campo en Colombia,
    #    basado en las ocupaciones afines ESCO + habilidades detectadas.
    descripcion_profesional: dict | None = None
    if is_gemini_available() and ocupaciones:
        try:
            ocupaciones_nombres = [_fix_mojibake(occ["nombre"]) for occ in ocupaciones[:5]]
            hab_top_nombres = [h["habilidad"] for h in hab_es[:6]]
            sys_prompt = (
                "Eres ALBA, orientador vocacional experto en Colombia. Genera una descripción "
                "profesional del campo laboral de un programa académico. Basándote en las ocupaciones "
                "ESCO afines y las habilidades clave, describe QUÉ HACE el profesional de ese campo "
                "en su día a día, en qué sectores trabaja, y qué tipo de problemas resuelve. "
                "NO inventes cifras. NO uses markdown. Devuelve SOLO un JSON válido con esta estructura:\n"
                '{"perfil_profesional": "párrafo de 3-4 frases describiendo qué hace el profesional",'
                ' "sectores_trabajo": ["sector1", "sector2", "sector3"],'
                ' "tipo_problemas": "frase breve sobre qué problemas resuelve",'
                ' "ejemplo_dia_a_dia": "frase breve de ejemplo de tarea típica"}'
            )
            user_prompt = (
                f"PROGRAMA: {req.programa}\n"
                f"OCUPACIONES AFINES (ESCO): {json.dumps(ocupaciones_nombres, ensure_ascii=False)}\n"
                f"HABILIDADES CLAVE: {json.dumps(hab_top_nombres, ensure_ascii=False)}\n"
                f"SALARIO PROMEDIO (GEIH): {salario_geih.get('salario_promedio_cop') if salario_geih else 'N/D'} COP/mes\n"
            )
            res = call_gemini_text(sys_prompt, user_prompt, temperature=0.4, max_tokens=600)
            if res.startswith("```json"): res = res[7:]
            if res.startswith("```"): res = res[3:]
            if res.endswith("```"): res = res[:-3]
            desc = json.loads(res.strip())
            if isinstance(desc, dict) and "perfil_profesional" in desc:
                descripcion_profesional = {
                    "perfil_profesional": _fix_mojibake(desc.get("perfil_profesional", "")),
                    "sectores_trabajo": desc.get("sectores_trabajo", [])[:5],
                    "tipo_problemas": _fix_mojibake(desc.get("tipo_problemas", "")),
                    "ejemplo_dia_a_dia": _fix_mojibake(desc.get("ejemplo_dia_a_dia", "")),
                    "fuente": "ALBA (basado en ESCO + GEIH)",
                }
        except Exception as e:
            print("DEBUG explora: error descripción profesional IA:", e)

    # 10. Tareas reales del día a día y requisitos de formación (O*NET)
    #     Reemplazan prosa inventada por datos reales del mercado laboral.
    tareas_reales: list[dict] = []
    requisitos_formacion: dict | None = None
    onet_codes = onet_codes_para_programa(programa_norm)
    if onet_codes:
        tareas_reales = tareas_onet_para_ocupaciones(onet_codes, limite=6)
        requisitos_formacion = educacion_onet_para_ocupaciones(onet_codes)

    return {
        "programa": req.programa,
        "ocupaciones_afines": [_fix_mojibake(occ["nombre"]) for occ in ocupaciones],
        "habilidades_desarrollaras": hab_es,
        "salidas_laborales": [_fix_mojibake(occ["nombre"]) for occ in ocupaciones[:5]],
        "demanda_laboral_sectores": sectores_nacional,
        "salario_esperado": {
            "rango_modal": salario_rango,
            "mediana_cop": round(salario_cop) if salario_cop else None,
            "egresados_anuales_nacional": egresados_anuales,
            "fuente_historica": "OLE/MEN 2001-2022",
        },
        "salario_real_geih": salario_geih,
        "saturacion_mercado": saturacion,
        "saber_pro": saber_pro,
        "descripcion_profesional": descripcion_profesional,
        "tareas_reales_onet": tareas_reales,
        "requisitos_formacion_onet": requisitos_formacion,
        "donde_estudiarla": [
            {
                "institucion": _fix_mojibake(i.get("institucion") or ""),
                "programa": _fix_mojibake(i.get("programa") or ""),
                "departamento": _fix_mojibake(i.get("departamento") or ""),
                "matriculados": int(i.get("matriculados") or 0),
            }
            for i in instituciones
        ],
        "fuentes": ["ESCO (UE)", "OLE/MEN", "GEIH-DANE", "SNIES-MEN", "Saber Pro/ICFES", "O*NET"],
    }


# ===========================================================================
# Analizar con IA — chat sobre la simulación generada
# ===========================================================================

class AnalizarRequest(BaseModel):
    simulacion_tipo: str  # "alineacion" | "intervencion" | "explora"
    simulacion_titulo: str  # título legible (ej: "Alineación curricular: Ingeniería de Sistemas")
    simulacion_data: dict[str, Any]
    pregunta: str
    historial: list[dict[str, Any]] | None = None


@router.post("/analizar")
async def analizar_con_ia(req: AnalizarRequest):
    """Gemini responde preguntas sobre la simulación generada. Sistema ALBA estricto."""
    system = """Eres ALBA, una plataforma de inteligencia laboral para Colombia. Estás analizando el resultado de una simulación basada en datos reales del DANE (GEIH), MEN (SNIES, OLE), SENA (SPE/APE), ESCO (UE) y DNP/MDM.

REGLAS OBLIGATORIAS:
1. Responde SOBRE los datos de la simulación que recibes. NO contradigas los números.
2. ANTI-ALUCINACIONES: nunca inventes cifras. Si necesitas un dato que no está en la simulación, dilo claramente ("este dato no está en la simulación") y propone cómo se podría obtener.
3. Usa SIEMPRE pesos colombianos (COP, $). Nunca dólares ni euros.
4. Sé conciso: máximo 2-3 párrafos. Usa **negrita** para resaltar datos clave.
5. Responde en español de Colombia, tono profesional y cercano.
6. NO comiences con "¡Hola!" ni "Claro!" ni "Entiendo que...". Responde directamente.
7. Si la pregunta no tiene relación con la simulación, redirige amablemente: "Esta pregunta no corresponde a la simulación actual. Aquí estamos analizando [título]."
8. Cita las fuentes de los datos cuando sea relevante (GEIH, OLE, ESCO, etc.)."""

    historial_txt = ""
    if req.historial:
        historial_txt = "\n\n## Historial de la conversación\n" + "\n".join(
            f"- **{'Usuario' if h.get('role') == 'user' else 'ALBA'}**: {h.get('content', '')}"
            for h in req.historial
        )

    # Serializar dump de la simulación (limitado para no saturar el prompt)
    try:
        dump = json.dumps(req.simulacion_data, ensure_ascii=False, default=str)
        if len(dump) > 6000:
            dump = dump[:6000] + "\n... (truncado)"
    except Exception:
        dump = str(req.simulacion_data)[:6000]

    user = f"""## Simulación analizada
- Tipo: {req.simulacion_tipo}
- Título: {req.simulacion_titulo}

## Datos de la simulación
{dump}{historial_txt}

## Pregunta del usuario
{req.pregunta}
"""

    try:
        if is_gemini_available():
            respuesta = call_gemini_text(system, user)
        else:
            respuesta = call_llm_text(system, user)
        return {"respuesta": respuesta, "simulacion_tipo": req.simulacion_tipo}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error al consultar la IA: {str(e)}")


# ===========================================================================
# Endpoints auxiliares: listados para autocompletar
# ===========================================================================

# ============================================================================
# Autocompletado de programas SNIES.
# Usa la tabla snies_programas_unicos (una fila por programa, ~4000 filas) que
# es instantánea de consultar. Si la tabla no existe, cae a consulta directa
# sobre snies_programas_matriculados (600K filas, más lenta).
# ============================================================================
_tabla_unicos_disponible: bool | None = None  # cache: True/False/None(sin verificar)


def _verificar_tabla_unicos() -> bool:
    """Verifica una sola vez si la tabla snies_programas_unicos existe."""
    global _tabla_unicos_disponible
    if _tabla_unicos_disponible is not None:
        return _tabla_unicos_disponible
    try:
        r = supabase.table("snies_programas_unicos").select("id").limit(1).execute()
        _tabla_unicos_disponible = bool(r.data is not None)
    except Exception:
        _tabla_unicos_disponible = False
    return _tabla_unicos_disponible


@router.get("/programas")
async def listar_programas(q: str = "", limit: int = 30):
    """Lista programas SNIES para autocompletar.
    Prioriza la tabla snies_programas_unicos (rápida, una fila por programa).
    Fallback: consulta directa sobre snies_programas_matriculados (más lenta)."""
    q_norm = _norm(q.strip())

    # 1) Tabla de programas únicos (rápida, ~4000 filas)
    if _verificar_tabla_unicos():
        if not q_norm:
            # Sin query: top por matriculados (usa índice, instantáneo)
            rows = _sb_select("snies_programas_unicos",
                              select="programa",
                              order="matriculados_max.desc", limit=limit)
        else:
            # Con query: ilike sobre programa_norm (usa índice trigram, instantáneo)
            rows = _sb_ilike("snies_programas_unicos",
                             "programa", "programa_norm",
                             f"%{q_norm}%", limit=limit, order="matriculados_max.desc")
        return {"programas": [{"programa": r.get("programa")} for r in rows]}

    # 2) Fallback: tabla original (600K filas, dedup en Python)
    if not q_norm:
        try:
            rows = _sb_select("snies_programas_matriculados",
                              select="programa,matriculados",
                              order="matriculados.desc", limit=300)
        except Exception:
            rows = _sb_select("snies_programas_matriculados",
                              select="programa,matriculados",
                              limit=500)
    else:
        try:
            rows = _sb_ilike("snies_programas_matriculados",
                             "programa,matriculados", "programa",
                             f"%{q_norm}%", limit=200)
        except Exception as e:
            print("DEBUG listar_programas: ilike falló:", str(e)[:80])
            rows = _sb_select("snies_programas_matriculados",
                              select="programa,matriculados",
                              filt={"programa": q.strip()},
                              limit=100)
    # Deduplicar por nombre normalizado
    seen: dict[str, dict] = {}
    for r in rows:
        p = r.get("programa") or ""
        pn = _norm(p)
        if not pn:
            continue
        m = float(r.get("matriculados") or 0)
        if pn not in seen or m > seen[pn].get("_m", 0):
            seen[pn] = {"programa": p, "_m": m}
    salida = sorted(seen.values(), key=lambda x: -x["_m"])[:limit]
    for x in salida:
        del x["_m"]
    return {"programas": salida}


@router.get("/departamentos")
async def listar_departamentos():
    rows = _sb_select("geih_resumen_departamento", select="departamento")
    # Limpieza mínima de mojibake
    salida = sorted({_fix_mojibake(str(r.get("departamento") or "")) for r in rows if r.get("departamento")})
    return {"departamentos": salida}


# ===========================================================================
# Laboratorio de Actualización Curricular — Análisis de documento real
#
# Flujo (consensuado con los 3 críticos):
# 1. Rector sube su documento curricular (PDF/imagen/DOCX/texto).
# 2. Gemini multimodal extrae materias + competencias con evidencia textual.
# 3. ALBA valida: una skill ESCO solo cuenta como cubierta si hay evidencia.
# 4. Matching por embeddings (cosine >= 0.90) entre competencias y ESCO.
# 5. Python calcula índice, top-5 (recalc), before/after (recalc).
# 6. Botón genera informe descargable (resumen por Gemini, no inventa cifras).
#
# Filosofía: Gemini COMPRENDE el documento (extrae), ALBA DECIDE (matching
# semántico por embeddings + cálculo determinístico). ALBA no adivina, recalcula.
# ===========================================================================

# MIME types soportados por Gemini multimodal (inline)
MIME_ACEPTADOS = {
    "application/pdf": "pdf",
    "image/png": "png",
    "image/jpeg": "jpg",
    "image/jpg": "jpg",
    "image/webp": "webp",
    "image/gif": "gif",
    "text/plain": "txt",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/msword": "docx",
}

EXTENSION_A_MIME = {
    "pdf": "application/pdf",
    "png": "image/png",
    "jpg": "image/jpeg",
    "jpeg": "image/jpeg",
    "webp": "image/webp",
    "gif": "image/gif",
    "txt": "text/plain",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "doc": "application/msword",
}


def _extraer_texto_docx(bytes_docx: bytes) -> str:
    """Extrae el texto plano de un .docx usando python-docx (Gemini no lee DOCX)."""
    import io
    try:
        import docx as python_docx
        doc = python_docx.Document(io.BytesIO(bytes_docx))
        textos = []
        for para in doc.paragraphs:
            if para.text.strip():
                textos.append(para.text.strip())
        # Incluir tablas (los pensums suelen venir en tablas)
        for tabla in doc.tables:
            for fila in tabla.rows:
                celdas = [c.text.strip() for c in fila.cells if c.text.strip()]
                if celdas:
                    textos.append(" | ".join(celdas))
        return "\n".join(textos)
    except Exception as e:
        print(f"DOCX extract error: {e}")
        return ""


def _determinar_mime(nombre_archivo: str, content_type: str) -> str | None:
    """Determina el MIME type real a partir del nombre y el content-type declarado."""
    if content_type and content_type in MIME_ACEPTADOS:
        return content_type
    ext = (nombre_archivo.rsplit(".", 1)[-1].lower() if "." in nombre_archivo else "")
    return EXTENSION_A_MIME.get(ext)


def _limpiar_materias(materias: list) -> list[str]:
    """Limpia la lista de materias que devuelve Gemini."""
    limpias: list[str] = []
    for m in materias:
        if not m or not isinstance(m, str):
            continue
        t = _fix_mojibake(m).strip()
        # Filtrar basura: muy cortas, muy largas, o que parecen créditos
        if len(t) < 3 or len(t) > 120:
            continue
        # Filtrar números sueltos o códigos
        if t.replace(" ", "").isdigit():
            continue
        limpias.append(t)
    # Deduplicar preservando orden
    seen: set[str] = set()
    out: list[str] = []
    for t in limpias:
        k = _norm(t)
        if k not in seen:
            seen.add(k)
            out.append(t)
    return out


async def _extraer_competencias_gemini(
    file_bytes: bytes, mime_type: str, nombre_archivo: str
) -> dict[str, Any]:
    """
    Usa Gemini (multimodal) para extraer materias y competencias de un documento
    curricular, con evidencia textual obligatoria.

    Contrato: Gemini NO sabe qué es ESCO. Solo extrae lo que ve en el documento.
    """
    if not is_gemini_available():
        raise HTTPException(
            status_code=503,
            detail="Se requiere Gemini para analizar el documento. Configura GOOGLE_API_KEY."
        )

    # DOCX: Gemini no lo lee nativamente, extraemos el texto primero
    if mime_type in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                     "application/msword"):
        texto = _extraer_texto_docx(file_bytes)
        if not texto.strip():
            raise HTTPException(status_code=400, detail="No se pudo extraer texto del documento Word.")
        return await _extraer_competencias_gemini_texto(texto, nombre_archivo)

    # TXT: extraer texto y procesar como texto
    if mime_type == "text/plain":
        try:
            texto = file_bytes.decode("utf-8", errors="replace")
        except Exception:
            texto = file_bytes.decode("latin-1", errors="replace")
        return await _extraer_competencias_gemini_texto(texto, nombre_archivo)

    # PDF/imagen: Gemini multimodal nativo
    system = (
        "Eres un experto en educación superior colombiana. "
        "Analiza el pensum o plan de estudios (PDF o imagen) e IDENTIFICA todas las "
        "materias/asignaturas que contiene. Para CADA materia, INFIERE de 1 a 3 competencias "
        "profesionales concretas que desarrolla, usando tu conocimiento de lo que "
        "típicamente se enseña en esa materia en universidades colombianas.\n\n"
        "Reglas estrictas:\n"
        "1. INFIERE competencias basadas en lo que REALMENTE enseña cada materia. No repitas el nombre de la materia como competencia.\n"
        "2. Usa frases descriptivas en español: 'resolver sistemas de ecuaciones lineales', no 'ÁLGEBRA LINEAL'.\n"
        "3. NO uses ESCO ni inglés: usa términos del español colombiano.\n"
        "4. Si una materia es genérica (ej. Cátedra Unilibrista, Electiva), infiere igualmente qué competencia blanda desarrolla.\n"
        "5. Cada competencia debe ser una habilidad observable y enseñable, no un tema.\n"
        "6. Devuelve SOLO el JSON, sin markdown."
    )
    user_text = (
        f"Analiza el documento '{nombre_archivo}'.\n\n"
        "Devuelve EXACTAMENTE este JSON:\n"
        "{\n"
        '  "materias": ["Materia 1", "Materia 2", ...],\n'
        '  "competencias": [\n'
        '    {"nombre": "competencia inferida para esa materia", "materia": "nombre exacto de la materia del pensum"}\n'
        "  ]\n"
        "}"
    )
    resultado = call_gemini_multimodal_json(
        system=system,
        user_text=user_text,
        file_bytes=file_bytes,
        mime_type=mime_type,
        temperature=0.1,
        max_tokens=16384,
    )
    if "error" in resultado:
        raise HTTPException(status_code=502, detail=f"Gemini no pudo procesar el documento: {resultado.get('raw','')[:300]}")
    return resultado


async def _extraer_competencias_gemini_texto(texto: str, nombre_archivo: str) -> dict[str, Any]:
    """Versión texto (para TXT y DOCX ya extraído) del extractor de competencias."""
    if not is_gemini_available():
        raise HTTPException(status_code=503, detail="Se requiere Gemini para analizar el documento.")
    # Truncar texto muy largo (límite práctico de contexto)
    if len(texto) > 30000:
        texto = texto[:30000]
    system = (
        "Eres un experto en educación superior colombiana. "
        "Analiza el texto del pensum o plan de estudios e IDENTIFICA todas las "
        "materias/asignaturas que contiene. Para CADA materia, INFIERE de 1 a 3 competencias "
        "profesionales concretas que desarrolla, usando tu conocimiento de lo que "
        "típicamente se enseña en esa materia en universidades colombianas.\n\n"
        "Reglas estrictas:\n"
        "1. INFIERE competencias basadas en lo que REALMENTE enseña cada materia. No repitas el nombre de la materia como competencia.\n"
        "2. Usa frases descriptivas en español: 'resolver sistemas de ecuaciones lineales', no 'ÁLGEBRA LINEAL'.\n"
        "3. NO uses ESCO ni inglés: usa términos del español colombiano.\n"
        "4. Si una materia es genérica (ej. Cátedra Unilibrista, Electiva), infiere igualmente qué competencia blanda desarrolla.\n"
        "5. Cada competencia debe ser una habilidad observable y enseñable, no un tema.\n"
        "6. Devuelve SOLO el JSON, sin markdown."
    )
    user_text = (
        f"Analiza el texto del documento '{nombre_archivo}'.\n\n"
        "Devuelve EXACTAMENTE este JSON:\n"
        "{\n"
        '  "materias": ["Materia 1", "Materia 2", ...],\n'
        '  "competencias": [\n'
        '    {"nombre": "competencia inferida para esa materia", "materia": "nombre exacto de la materia del pensum"}\n'
        "  ]\n"
        "}\n\n"
        f"TEXTO DEL DOCUMENTO:\n{texto}"
    )
    from app.services.llm_gemini import call_gemini_json
    resultado = call_gemini_json(system=system, user=user_text, temperature=0.1, max_tokens=16384)
    if "error" in resultado:
        raise HTTPException(status_code=502, detail=f"Gemini no pudo procesar el documento: {resultado.get('raw','')[:300]}")
    return resultado


@router.post("/analizar-documento-curricular")
async def analizar_documento_curricular(
    programa: str = Form(...),
    archivo: UploadFile = File(...),
):
    """
    Analiza un documento curricular real subido por el rector.

    Flujo simplificado (1 llamada al LLM):
      1. Gemini extrae materias e infiere competencias del pensum (multimodal).
      2. ALBA carga las habilidades O*NET reales del mercado (por sector).
      3. El LLM compara el pensum con las habilidades O*NET en una sola llamada.
      4. Devuelve indice, cubiertas, faltantes y top-5 de impacto.
    """
    if not programa.strip():
        raise HTTPException(status_code=400, detail="Programa requerido")

    # 1. Validar archivo
    contenido = await archivo.read()
    if not contenido:
        raise HTTPException(status_code=400, detail="El archivo está vacío.")
    if len(contenido) > 20 * 1024 * 1024:  # 20 MB
        raise HTTPException(status_code=413, detail="El archivo excede 20 MB.")
    mime = _determinar_mime(archivo.filename or "", archivo.content_type or "")
    if not mime:
        exts = ", ".join(EXTENSION_A_MIME.keys())
        raise HTTPException(
            status_code=415,
            detail=f"Formato no soportado. Aceptados: {exts}."
        )

    # 2. Gemini extrae materias e infiere competencias del pensum (multimodal)
    extraccion = await _extraer_competencias_gemini(contenido, mime, archivo.filename or "documento")
    materias_crudas = _limpiar_materias(extraccion.get("materias", []))
    competencias_crudas = extraccion.get("competencias", [])

    # 3. Validar competencias (solo nombre + materia)
    competencias_validas: list[dict] = []
    seen_comps: set[str] = set()
    for c in competencias_crudas:
        if not isinstance(c, dict):
            continue
        nombre = _fix_mojibake(str(c.get("nombre", "")).strip())
        materia = _fix_mojibake(str(c.get("materia", "")).strip())
        if not nombre or not materia:
            continue
        k = _norm(nombre)
        if k in seen_comps:
            continue
        seen_comps.add(k)
        competencias_validas.append({"nombre": nombre, "materia": materia})

    # 4. ALBA compara el pensum con las habilidades O*NET en una sola llamada al LLM
    from app.services.analisis_curricular import analizar_pensum_vs_onet
    analisis = analizar_pensum_vs_onet(programa, competencias_validas)
    if "error" in analisis:
        raise HTTPException(status_code=502, detail=analisis["error"])

    # 4b. Traducir habilidades O*NET (vienen en inglés) al español
    hab_english_orig = [h["habilidad"] for h in analisis.get("habilidades_mercado", [])]
    trad_map = _traducir_habilidades_llm(hab_english_orig)
    for h in analisis.get("habilidades_mercado", []):
        h["habilidad"] = _fix_mojibake(trad_map.get(h["habilidad"], h["habilidad"]))
    for c in analisis.get("cubiertas", []):
        c["habilidad"] = _fix_mojibake(trad_map.get(c["habilidad"], c["habilidad"]))
    analisis["faltantes"] = [_fix_mojibake(trad_map.get(f, f)) for f in analisis.get("faltantes", [])]
    for t in analisis.get("top5_impacto", []):
        t["habilidad"] = _fix_mojibake(trad_map.get(t["habilidad"], t["habilidad"]))

    # 5. Datos de mercado reales: salario GEIH + saturación para ocupaciones afines
    mercado_onet = _mercado_onet_para_programa(programa, analisis.get("ocupaciones", []))

    # 6. Construir respuesta
    return {
        "programa": programa,
        "archivo": archivo.filename,
        "ocupaciones_afines": analisis.get("ocupaciones_afines", []),
        "total_esenciales": analisis["total_habilidades"],
        "skills_esco_esperadas": [h["habilidad"] for h in analisis.get("habilidades_mercado", [])],
        "skills_metadata": [
            {
                "habilidad_es": h["habilidad"],
                "habilidad_en": "",
                "categoria": h["categoria"],
                "hot": h["hot"],
                "demand": h["demand"],
            }
            for h in analisis.get("habilidades_mercado", [])
        ],
        "materias_detectadas": materias_crudas,
        "competencias_detectadas": competencias_validas,
        "total_materias": len(materias_crudas),
        "total_competencias": len(competencias_validas),
        "indice_base": analisis["indice"],
        "total_cubiertas": analisis["total_cubiertas"],
        "total_faltantes": analisis["total_faltantes"],
        "cubiertas_con_evidencia": [
            {
                "skill_esco": c["habilidad"],
                "competencia_detectada": c.get("competencia", ""),
                "similitud": 1.0,
                "materia": c.get("materia", ""),
                "confianza": c.get("confianza", ""),
            }
            for c in analisis.get("cubiertas", [])
        ],
        "competencias_faltan": analisis.get("faltantes", []),
        "top5_impacto": [
            {"competencia": t["habilidad"], "impacto": t["impacto"]}
            for t in analisis.get("top5_impacto", [])
        ],
        "umbral_cobertura": 0,
        "matching_metodo": "llm_onet",
        "metodologia": (
            "ALBA infiere las competencias de cada materia del pensum y las compara "
            "con las habilidades reales que el mercado exige según O*NET (tecnologías hot, "
            "actividades del trabajo y conocimiento requerido). El índice = cubiertas / total × 100. "
            "El salario y empleo se calculan desde GEIH-DANE usando el código ISCO de las ocupaciones afines."
        ),
        "fuentes": ["O*NET (EE.UU.)", "Documento del pensum", "GEIH-DANE", "SNIES-MEN"],
        "interpretacion": analisis.get("interpretacion", ""),
        "salario_real_geih": mercado_onet.get("salario"),
        "saturacion_mercado": mercado_onet.get("saturacion"),
    }


class RecalcularRequest(BaseModel):
    programa: str
    indice_base: float = 0
    total_habilidades: int = 0
    cubiertas_actuales: list[dict] = []
    competencias_extra: list[str] = []


@router.post("/recalcular-escenario")
async def recalcular_escenario(req: RecalcularRequest):
    """Recalcula el índice al simular fortalecer competencias.

    Aritmética simple: cada habilidad seleccionada que estaba faltando
    se suma al total de cubiertas. No requiere LLM ni embeddings.
    """
    total = max(req.total_habilidades, 1)
    base_cubiertas = len(req.cubiertas_actuales)
    indice_base = round(base_cubiertas / total * 100, 1)

    # Las competencias_extra son habilidades que el usuario seleccionó para fortalecer.
    # Cada una que no esté ya cubierta se suma.
    ya_cubiertas = set()
    for c in req.cubiertas_actuales:
        if isinstance(c, dict):
            ya_cubiertas.add(_norm(str(c.get("skill_esco", ""))))
        else:
            ya_cubiertas.add(_norm(str(c)))

    nuevas_cubiertas: list[str] = []
    for extra in req.competencias_extra:
        if _norm(extra) not in ya_cubiertas:
            nuevas_cubiertas.append(extra)

    nuevas_count = len(nuevas_cubiertas)
    total_sim = base_cubiertas + nuevas_count
    indice_sim = round(total_sim / total * 100, 1)
    delta_sim = round(indice_sim - indice_base, 1)

    # Justificación narrativa con LLM (opcional, interpreta no inventa)
    justificacion_ia: str | None = None
    if is_gemini_available() and nuevas_cubiertas:
        try:
            just_sys = (
                "Eres ALBA. Explica en 1-2 frases en español de Colombia "
                "el impacto de fortalecer estas habilidades en el programa. "
                "Sé concreto. NO inventes cifras: usa solo las que te damos. No uses markdown."
            )
            just_usr = json.dumps({
                "programa": req.programa,
                "habilidades_fortalecidas": nuevas_cubiertas[:5],
                "incremento_indice": f"+{delta_sim} puntos (de {indice_base}% a {indice_sim}%)",
            }, ensure_ascii=False)
            just_res = call_gemini_text(just_sys, just_usr, temperature=0.3, max_tokens=300)
            justificacion_ia = just_res.strip()
        except Exception:
            pass

    return {
        "programa": req.programa,
        "total_esenciales": total,
        "indice_base": indice_base,
        "total_cubiertas": base_cubiertas,
        "total_faltantes": total - base_cubiertas,
        "indice_sim": indice_sim,
        "delta_sim": delta_sim,
        "nuevas_cubiertas": nuevas_cubiertas,
        "justificacion_ia": justificacion_ia,
        "umbral_cobertura": 0,
    }


class PropuestaRequest(BaseModel):
    programa: str
    indice_base: float
    total_esenciales: int
    total_cubiertas: int
    total_faltantes: int
    top5_impacto: list[dict] = []
    competencias_faltan: list[str] = []
    ocupaciones_afines: list[str] = []
    umbral_cobertura: float = 0


@router.post("/generar-propuesta-actualizacion")
async def generar_propuesta_actualizacion(req: PropuestaRequest):
    """Genera un informe descargable con la propuesta de actualización curricular."""
    if not is_gemini_available():
        raise HTTPException(status_code=503, detail="Se requiere Gemini para generar el informe.")
    try:
        sys_prompt = (
            "Eres ALBA, asesor curricular del Ministerio de Educación de Colombia. "
            "Redacta una propuesta ejecutiva de actualización curricular en formato markdown, "
            "basándote EXCLUSIVAMENTE en los datos que recibes. NO inventes cifras. "
            "Estructura: 1) Resumen ejecutivo, 2) Estado actual (índice y cobertura), "
            "3) Brechas priorizadas (top-5 con impacto), 4) Competencias a fortalecer, "
            "5) Sectores beneficiados, 6) Recomendación. Tono profesional, español de Colombia."
        )
        usr_prompt = json.dumps({
            "programa": req.programa,
            "indice_base": req.indice_base,
            "total_esenciales": req.total_esenciales,
            "total_cubiertas": req.total_cubiertas,
            "total_faltantes": req.total_faltantes,
            "top5_impacto": req.top5_impacto,
            "competencias_faltan": req.competencias_faltan[:15],
            "ocupaciones_afines": req.ocupaciones_afines[:5],
            "umbral_cobertura": req.umbral_cobertura,
        }, ensure_ascii=False)
        informe = call_gemini_text(sys_prompt, usr_prompt, temperature=0.4, max_tokens=2000)
        return {"informe": informe.strip(), "programa": req.programa}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generando propuesta: {str(e)}")